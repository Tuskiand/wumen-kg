from __future__ import annotations

import csv
import importlib
import io
import json
import math
import threading
import time
import zipfile
from collections import Counter
from itertools import combinations

import networkx as nx
import numpy as np
from sklearn.decomposition import PCA

from app.schemas.graph import (
    CompareNode,
    CompareNodeGroup,
    GraphNode,
    GraphSnapshot,
    NodeRankingGroup,
    PhysicianCategoryScores,
    PhysicianDoctorScore,
    PhysicianEmbeddingScatterPoint,
    PhysicianFeatureEmbedding,
    PhysicianFeatureSimilarityCandidate,
    PhysicianPathChain,
    PhysicianPathCompareResponse,
    PhysicianPathCompareSummary,
    PhysicianPathCompleteness,
    PhysicianPathEmbeddingProfile,
    PhysicianPathProfile,
    PhysicianPathSimilarityPair,
    PhysicianNodeCompareResponse,
    PhysicianNodeCompareSummary,
    PhysicianNodeProfile,
    PhysicianNodeRwrResult,
    PhysicianSubgraphAuditEdge,
    PhysicianSubgraphAuditNode,
    PhysicianSubgraphCompareResponse,
    PhysicianSubgraphCompareSummary,
    PhysicianSubgraphEdge,
    PhysicianSubgraphEmbeddingProfile,
    PhysicianSubgraphProfile,
    PhysicianSubgraphRelationStat,
    PhysicianSubgraphSimilarityPair,
    PhysicianSimilarityGroup,
    PhysicianSimilarityPair,
    RankedGraphNode,
    SharedCompareNode,
    SharedCompareNodeGroup,
    SharedPhysicianPath,
    SharedPhysicianSubgraphEdge,
)


class PhysicianCompareService:
    ANALYSIS_TYPES = {"A医家", "B病名", "C证型", "D病因", "E病机"}
    RESULT_TYPES = {"C证型": "patterns", "D病因": "causes", "E病机": "mechanisms"}
    PATH_DIRECT_TYPES = {"C证型", "D病因", "E病机"}
    PATH_RELATION_TYPES = {"C证型-D病因", "C证型-E病机", "D病因-E病机"}
    SUBGRAPH_FEATURE_TYPES = {"C证型", "D病因", "E病机"}
    SUBGRAPH_CATEGORY_ORDER = ("patterns", "causes", "mechanisms")
    SUBGRAPH_VECTOR_METRICS = ("Graph2Vec",)
    RANDOM_SEED = 42
    PRIMARY_RESTART_PROBABILITY = 0.25
    RWR_MAX_ITERATIONS = 100
    RWR_TOLERANCE = 1e-6
    TOP_K = 10
    _word2vec_class: type[object] | None = None
    _word2vec_import_lock = threading.Lock()

    @classmethod
    def _load_word2vec_class(cls) -> type[object]:
        if cls._word2vec_class is not None:
            return cls._word2vec_class
        with cls._word2vec_import_lock:
            if cls._word2vec_class is not None:
                return cls._word2vec_class
            try:
                module = importlib.import_module("gensim.models.word2vec")
            except ModuleNotFoundError as exc:
                raise ValueError("路径嵌入依赖 gensim，请先进入 KG 环境再运行。") from exc
            cls._word2vec_class = module.Word2Vec
            return cls._word2vec_class

    def _build_graph(self, snapshot: GraphSnapshot) -> nx.Graph:
        graph = nx.Graph()
        node_by_id = {node.id: node for node in snapshot.nodes if node.type in self.ANALYSIS_TYPES}
        for node in node_by_id.values():
            graph.add_node(node.id, data=node)
        for edge in snapshot.edges:
            if edge.source in node_by_id and edge.target in node_by_id and edge.source != edge.target:
                graph.add_edge(edge.source, edge.target)
        return graph

    @staticmethod
    def _nodes_from_graph(graph: nx.Graph) -> dict[str, GraphNode]:
        return {node_id: graph.nodes[node_id]["data"] for node_id in graph.nodes}

    @staticmethod
    def _doctor_node(snapshot: GraphSnapshot) -> GraphNode | None:
        return next((node for node in snapshot.nodes if node.type == "A医家"), None)

    @staticmethod
    def _disease_node(snapshot: GraphSnapshot, disease: str) -> GraphNode | None:
        return next((node for node in snapshot.nodes if node.type == "B病名" and node.name == disease), None)

    def compare_nodes(
        self,
        disease: str,
        contexts: list[dict[str, object]],
        similarity_rows: dict[str, list[dict[str, object]]],
        fastrp_payload: dict[str, object],
    ) -> PhysicianNodeCompareResponse:
        doctors = [context["doctor"] for context in contexts if isinstance(context.get("doctor"), GraphNode)]
        if not doctors:
            return PhysicianNodeCompareResponse(
                disease=disease,
                doctor_count=0,
                doctors=[],
                similarity=PhysicianSimilarityGroup(),
                fastrp_similarity=PhysicianSimilarityGroup(),
                shared_nodes=SharedCompareNodeGroup(),
                doctor_profiles=[],
                rwr=[],
                doctor_feature_embeddings=[],
                similarity_overview=[],
                feature_similarity_candidates=[],
                embedding_points=[],
                summary=PhysicianNodeCompareSummary(
                    primary_similarity_metric="Jaccard",
                    primary_embedding_metric="FastRP cosine",
                    shared_node_count=0,
                    pairwise_comparison_count=0,
                    primary_restart_probability=self.PRIMARY_RESTART_PROBABILITY,
                    message=f"当前病名“{disease}”没有可比较的医家数据。",
                ),
            )

        doctor_node_maps = self._collect_doctor_node_maps(contexts)
        shared_nodes = self._build_shared_nodes(doctor_node_maps)
        doctor_profiles = self._build_doctor_profiles(contexts, doctor_node_maps, shared_nodes)
        similarity = self._build_similarity_group(doctors, similarity_rows)
        fastrp_similarity = self._build_similarity_group(
            doctors,
            self._typed_mapping(fastrp_payload.get("similarity")),
        )
        rwr = self._build_rwr_results(disease, contexts)
        doctor_feature_embeddings = self._build_doctor_feature_embeddings(doctors, fastrp_payload)
        similarity_overview = self._build_doctor_score_overview(doctors, fastrp_similarity)
        feature_similarity_candidates = self._build_feature_similarity_candidates(fastrp_payload)
        embedding_points = self._build_embedding_points(self._typed_list(fastrp_payload.get("feature_points")))
        shared_node_count = sum(
            len(getattr(shared_nodes, category))
            for category in ("patterns", "causes", "mechanisms")
        )
        summary_message = self._node_compare_summary_message(disease, doctors, shared_node_count)
        return PhysicianNodeCompareResponse(
            disease=disease,
            doctor_count=len(doctors),
            doctors=doctors,
            similarity=similarity,
            fastrp_similarity=fastrp_similarity,
            shared_nodes=shared_nodes,
            doctor_profiles=doctor_profiles,
            rwr=rwr,
            doctor_feature_embeddings=doctor_feature_embeddings,
            similarity_overview=similarity_overview,
            feature_similarity_candidates=feature_similarity_candidates,
            embedding_points=embedding_points,
            summary=PhysicianNodeCompareSummary(
                primary_similarity_metric="Jaccard",
                primary_embedding_metric="FastRP cosine",
                shared_node_count=shared_node_count,
                pairwise_comparison_count=len(similarity.overall),
                primary_restart_probability=self.PRIMARY_RESTART_PROBABILITY,
                message=summary_message,
            ),
        )

    def compare_paths(
        self,
        disease: str,
        contexts: list[dict[str, object]],
    ) -> PhysicianPathCompareResponse:
        doctors = [context["doctor"] for context in contexts if isinstance(context.get("doctor"), GraphNode)]
        if not doctors:
            return PhysicianPathCompareResponse(
                disease=disease,
                doctor_count=0,
                doctors=[],
                shared_paths=[],
                doctor_profiles=[],
                similarity_pairs=[],
                embeddings=[],
                embedding_points=[],
                summary=PhysicianPathCompareSummary(
                    shared_path_count=0,
                    pairwise_comparison_count=0,
                    primary_similarity_metric="Path Jaccard",
                    embedding_metric="Metapath2Vec cosine",
                    message=f"当前病名“{disease}”没有可比较的辨证路径数据。",
                ),
            )

        doctor_paths = self._build_doctor_paths(contexts)
        shared_paths = self._build_shared_paths(doctor_paths)
        profiles = self._build_path_profiles(contexts, doctor_paths, shared_paths)
        embeddings = self._build_path_embedding_profiles(doctors, doctor_paths)
        similarity_pairs = self._build_path_similarity_pairs(doctors, doctor_paths, embeddings)
        embedding_points = self._build_doctor_embedding_points(embeddings, "路径向量")
        return PhysicianPathCompareResponse(
            disease=disease,
            doctor_count=len(doctors),
            doctors=doctors,
            shared_paths=shared_paths,
            doctor_profiles=profiles,
            similarity_pairs=similarity_pairs,
            embeddings=embeddings,
            embedding_points=embedding_points,
            summary=PhysicianPathCompareSummary(
                shared_path_count=len(shared_paths),
                pairwise_comparison_count=len(similarity_pairs),
                primary_similarity_metric="Path Jaccard",
                embedding_metric="Metapath2Vec cosine",
                message=self._path_compare_summary_message(disease, doctors, len(shared_paths)),
            ),
        )

    def compare_subgraphs(
        self,
        disease: str,
        contexts: list[dict[str, object]],
    ) -> PhysicianSubgraphCompareResponse:
        # 子图比较总流程：
        # 1. 为每位医家抽取“医家-病名核心子图”
        # 2. 汇总共同节点、共同边、独有节点、独有边
        # 3. 计算两两相似度：Jaccard + Graph2Vec 余弦相似度
        subgraphs = self._build_doctor_subgraphs(disease, contexts)
        doctors = [context["doctor"] for context in contexts if isinstance(context.get("doctor"), GraphNode) and context["doctor"].name in subgraphs]
        if not doctors:
            return PhysicianSubgraphCompareResponse(
                disease=disease,
                doctor_count=0,
                doctors=[],
                similarity_pairs=[],
                shared_nodes=SharedCompareNodeGroup(),
                shared_edges=[],
                doctor_profiles=[],
                embeddings=[],
                embedding_points=[],
                summary=PhysicianSubgraphCompareSummary(
                    primary_similarity_metric="子图Jaccard，辅以 Graph2Vec",
                    shared_node_count=0,
                    shared_edge_count=0,
                    pairwise_comparison_count=0,
                    vector_similarity_metrics=list(self.SUBGRAPH_VECTOR_METRICS),
                    message=f"当前病名“{disease}”没有可比较的核心子图数据。",
                ),
            )

        doctor_node_maps = {
            doctor_name: payload["node_map"]
            for doctor_name, payload in subgraphs.items()
            if isinstance(payload.get("node_map"), dict)
        }
        shared_nodes = self._build_shared_nodes(doctor_node_maps)
        shared_edges = self._build_shared_subgraph_edges(subgraphs)
        profiles = self._build_subgraph_profiles(contexts, subgraphs, shared_nodes, shared_edges)
        embedding_payload = self._build_subgraph_embeddings(sorted({doctor.name for doctor in doctors}), subgraphs)
        similarity_pairs = self._build_subgraph_similarity_pairs(doctors, subgraphs, embedding_payload)
        embeddings = self._build_subgraph_embedding_profiles(doctors, embedding_payload)
        embedding_points = self._build_doctor_embedding_points(embeddings, "Graph2Vec", vector_attr="graph2vec_vector")
        shared_node_count = sum(len(getattr(shared_nodes, category)) for category in self.SUBGRAPH_CATEGORY_ORDER)
        return PhysicianSubgraphCompareResponse(
            disease=disease,
            doctor_count=len(doctors),
            doctors=doctors,
            similarity_pairs=similarity_pairs,
            shared_nodes=shared_nodes,
            shared_edges=shared_edges,
            doctor_profiles=profiles,
            embeddings=embeddings,
            embedding_points=embedding_points,
            summary=PhysicianSubgraphCompareSummary(
                primary_similarity_metric="子图Jaccard，辅以 Graph2Vec",
                shared_node_count=shared_node_count,
                shared_edge_count=len(shared_edges),
                pairwise_comparison_count=len(similarity_pairs),
                vector_similarity_metrics=list(self.SUBGRAPH_VECTOR_METRICS),
                message=self._subgraph_compare_summary_message(disease, doctors, shared_node_count, len(shared_edges)),
            ),
        )

    def export_subgraphs(
        self,
        disease: str,
        contexts: list[dict[str, object]],
        similarity_rows: dict[str, list[dict[str, object]]],
        fastrp_payload: dict[str, object],
    ) -> tuple[str, bytes]:
        subgraphs = self._build_doctor_subgraphs(disease, contexts)
        doctors = [context["doctor"] for context in contexts if isinstance(context.get("doctor"), GraphNode) and context["doctor"].name in subgraphs]
        if not doctors:
            raise ValueError(f"当前病名“{disease}”没有可导出的核心子图数据。")

        doctor_node_maps = {
            doctor_name: payload["node_map"]
            for doctor_name, payload in subgraphs.items()
            if isinstance(payload.get("node_map"), dict)
        }
        shared_nodes = self._build_shared_nodes(doctor_node_maps)
        shared_edges = self._build_shared_subgraph_edges(subgraphs)
        profiles = self._build_subgraph_profiles(contexts, subgraphs, shared_nodes, shared_edges)
        embedding_payload = self._build_subgraph_embeddings(sorted({doctor.name for doctor in doctors}), subgraphs)
        similarity_pairs = self._build_subgraph_similarity_pairs(doctors, subgraphs, embedding_payload)
        archive = io.BytesIO()
        with zipfile.ZipFile(archive, "w", compression=zipfile.ZIP_DEFLATED) as bundle:
            bundle.writestr("README.txt", self._subgraph_export_readme(disease, doctors))
            bundle.writestr("summary.json", self._json_text({
                "disease": disease,
                "doctor_count": len(doctors),
                "doctors": [doctor.name for doctor in doctors],
                "vector_metrics": list(self.SUBGRAPH_VECTOR_METRICS),
            }))
            bundle.writestr("subgraph_jaccard_similarity.csv", self._csv_text(
                [
                    {
                        "left_doctor": pair.left_doctor,
                        "right_doctor": pair.right_doctor,
                        "node_jaccard": pair.node_jaccard,
                        "edge_jaccard": pair.edge_jaccard,
                        "subgraph_jaccard": pair.subgraph_jaccard,
                    }
                    for pair in similarity_pairs
                ],
                ["left_doctor", "right_doctor", "node_jaccard", "edge_jaccard", "subgraph_jaccard"],
            ))
            bundle.writestr("graph2vec_similarity.csv", self._csv_text(
                [
                    {
                        "left_doctor": pair.left_doctor,
                        "right_doctor": pair.right_doctor,
                        "graph2vec_cosine": pair.graph2vec_cosine,
                    }
                    for pair in similarity_pairs
                ],
                ["left_doctor", "right_doctor", "graph2vec_cosine"],
            ))
            bundle.writestr("pairwise_similarity.csv", self._csv_text(
                [
                    {
                        "left_doctor": pair.left_doctor,
                        "right_doctor": pair.right_doctor,
                        "node_jaccard": pair.node_jaccard,
                        "edge_jaccard": pair.edge_jaccard,
                        "subgraph_jaccard": pair.subgraph_jaccard,
                        "graph2vec_cosine": pair.graph2vec_cosine,
                    }
                    for pair in similarity_pairs
                ],
                [
                    "left_doctor",
                    "right_doctor",
                    "node_jaccard",
                    "edge_jaccard",
                    "subgraph_jaccard",
                    "graph2vec_cosine",
                ],
            ))
            bundle.writestr("shared_nodes.csv", self._csv_text(
                self._shared_node_rows(shared_nodes),
                ["category", "node_name", "node_type", "doctors"],
            ))
            bundle.writestr("shared_edges.csv", self._csv_text(
                [
                    {
                        "relation_type": item.edge.relation_type,
                        "edge_text": item.edge.text,
                        "doctors": "、".join(item.doctors),
                    }
                    for item in shared_edges
                ],
                ["relation_type", "edge_text", "doctors"],
            ))
            node_export = self._build_node_export_payload(disease, contexts, similarity_rows, fastrp_payload)
            bundle.writestr("node_fastrp_similarity.csv", self._csv_text(
                self._similarity_export_rows(node_export["similarity_group"]),
                ["category", "left_doctor", "right_doctor", "jaccard", "overlap", "cosine"],
            ))
            bundle.writestr("doctor_feature_embedding.csv", self._csv_text(
                self._doctor_feature_embedding_rows(node_export["embeddings"]),
                ["doctor", "category", "vector"],
            ))
            path_export = self._build_path_export_payload(disease, contexts)
            bundle.writestr("path_pairwise_similarity.csv", self._csv_text(
                [
                    {
                        "left_doctor": pair.left_doctor,
                        "right_doctor": pair.right_doctor,
                        "shared_path_count": pair.shared_path_count,
                        "union_path_count": pair.union_path_count,
                        "path_jaccard": pair.path_jaccard,
                        "metapath2vec_cosine": pair.metapath2vec_cosine,
                    }
                    for pair in path_export["similarity_pairs"]
                ],
                ["left_doctor", "right_doctor", "shared_path_count", "union_path_count", "path_jaccard", "metapath2vec_cosine"],
            ))
            bundle.writestr("shared_paths.csv", self._csv_text(
                self._shared_path_rows(path_export["shared_paths"]),
                ["path_type", "path_text", "doctors"],
            ))
            bundle.writestr("unique_paths.csv", self._csv_text(
                self._unique_path_rows(path_export["profiles"]),
                ["doctor", "path_type", "path_text", "path_category"],
            ))
            bundle.writestr("subgraph_method_parameters.json", self._json_text(
                {
                    "random_seed": self.RANDOM_SEED,
                    "subgraph_vector_metrics": list(self.SUBGRAPH_VECTOR_METRICS),
                    "path_embedding": "Metapath2Vec",
                    "node_embedding": "FastRP",
                }
            ))

            profile_by_doctor = {profile.doctor.name: profile for profile in profiles}
            path_profile_by_doctor = {profile.doctor.name: profile for profile in path_export["profiles"]}
            path_embedding_by_doctor = {item.doctor.name: item for item in path_export["embeddings"]}
            for doctor in doctors:
                doctor_name = doctor.name
                safe_doctor = self._safe_filename_component(doctor_name)
                payload = subgraphs.get(doctor_name, {})
                profile = profile_by_doctor.get(doctor_name)
                if not isinstance(payload, dict) or profile is None:
                    continue
                prefix = f"doctors/{safe_doctor}"
                bundle.writestr(f"{prefix}/core_nodes.csv", self._csv_text(
                    self._core_node_rows(doctor, disease, payload),
                    ["id", "name", "type", "role", "category", "inclusion_reason"],
                ))
                bundle.writestr(f"{prefix}/core_edges.csv", self._csv_text(
                    [
                        {
                            "relation_type": edge.relation_type,
                            "source_name": edge.source_name,
                            "source_type": edge.source_type,
                            "target_name": edge.target_name,
                            "target_type": edge.target_type,
                            "text": edge.text,
                        }
                        for edge in profile.edges
                    ],
                    ["relation_type", "source_name", "source_type", "target_name", "target_type", "text"],
                ))
                bundle.writestr(f"{prefix}/audit_nodes.csv", self._csv_text(
                    [
                        {
                            "id": node.id,
                            "name": node.name,
                            "type": node.type,
                            "inclusion_reason": node.inclusion_reason,
                        }
                        for node in profile.audit_nodes
                    ],
                    ["id", "name", "type", "inclusion_reason"],
                ))
                bundle.writestr(f"{prefix}/audit_edges.csv", self._csv_text(
                    [
                        {
                            "relation_type": edge.relation_type,
                            "source_name": edge.source_name,
                            "source_type": edge.source_type,
                            "target_name": edge.target_name,
                            "target_type": edge.target_type,
                            "text": edge.text,
                        }
                        for edge in profile.audit_edges
                    ],
                    ["relation_type", "source_name", "source_type", "target_name", "target_type", "text"],
                ))
                vector_graph = payload.get("vector_graph")
                if isinstance(vector_graph, nx.Graph):
                    bundle.writestr(f"{prefix}/relation_node_graph.json", self._json_text(self._serialize_vector_graph(vector_graph)))
                graph2vec_vector = embedding_payload["vector_embeddings"].get("Graph2Vec", {}).get(doctor_name)
                if graph2vec_vector is not None:
                    bundle.writestr(f"{prefix}/graph2vec_vector.csv", self._csv_text(self._vector_rows(graph2vec_vector), ["index", "value"]))
                path_profile = path_profile_by_doctor.get(doctor_name)
                if path_profile is not None:
                    bundle.writestr(f"{prefix}/diagnostic_paths.csv", self._csv_text(
                        self._doctor_path_rows(path_profile),
                        ["path_category", "path_type", "path_text", "signature"],
                    ))
                    bundle.writestr(f"{prefix}/path_statistics.csv", self._csv_text(
                        [self._path_statistics_row(path_profile)],
                        ["doctor", "complete_count", "partial_count", "single_count", "total_count", "complete_ratio", "path_coverage"],
                    ))
                path_embedding = path_embedding_by_doctor.get(doctor_name)
                if path_embedding is not None:
                    bundle.writestr(f"{prefix}/metapath2vec_path_vector.csv", self._csv_text(
                        self._vector_rows(np.asarray(path_embedding.vector, dtype=float)),
                        ["index", "value"],
                    ))

        filename = f"physician_subgraphs_{self._safe_filename_component(disease)}.zip"
        return filename, archive.getvalue()

    def export_paper_report(
        self,
        disease: str,
        contexts: list[dict[str, object]],
        similarity_rows: dict[str, list[dict[str, object]]],
        fastrp_payload: dict[str, object],
    ) -> tuple[str, bytes, dict[str, int]]:
        total_start = time.perf_counter()
        node_result = self.compare_nodes(disease, contexts, similarity_rows, fastrp_payload)
        path_result = self.compare_paths(disease, contexts)
        subgraph_result = self.compare_subgraphs(disease, contexts)
        if not node_result.doctors:
            raise ValueError(f"当前病名“{disease}”没有可导出的医家比较结果。")
        analysis_ms = self._elapsed_ms(total_start)

        figure_start = time.perf_counter()
        figures = self._build_report_figures(disease, node_result, path_result, subgraph_result)
        figure_ms = self._elapsed_ms(figure_start)

        table_payload = self._build_report_tables(node_result, path_result, subgraph_result)

        word_start = time.perf_counter()
        report_docx = self._build_report_docx(
            disease,
            node_result,
            path_result,
            subgraph_result,
            figures,
            table_payload,
            {"analysis_ms": analysis_ms, "figure_ms": figure_ms, "word_ms": 0, "total_ms": 0},
        )
        word_ms = self._elapsed_ms(word_start)
        final_timings = {
            "analysis_ms": analysis_ms,
            "figure_ms": figure_ms,
            "word_ms": word_ms,
            "total_ms": self._elapsed_ms(total_start),
        }
        table_payload["timings.csv"] = [
            {"stage": "analysis", "elapsed_ms": final_timings["analysis_ms"]},
            {"stage": "figure_generation", "elapsed_ms": final_timings["figure_ms"]},
            {"stage": "word_generation", "elapsed_ms": final_timings["word_ms"]},
            {"stage": "total", "elapsed_ms": final_timings["total_ms"]},
        ]
        report_docx = self._build_report_docx(
            disease,
            node_result,
            path_result,
            subgraph_result,
            figures,
            table_payload,
            final_timings,
        )
        final_timings["total_ms"] = self._elapsed_ms(total_start)

        metadata = {
            "disease": disease,
            "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "doctor_count": len(node_result.doctors),
            "figures": sorted(figures.keys()),
            "tables": sorted(table_payload.keys()),
            "timings_ms": {
                "analysis": final_timings["analysis_ms"],
                "figure_generation": final_timings["figure_ms"],
                "word_generation": final_timings["word_ms"],
                "total": final_timings["total_ms"],
            },
        }

        archive = io.BytesIO()
        safe_disease = self._safe_filename_component(disease)
        with zipfile.ZipFile(archive, "w", compression=zipfile.ZIP_DEFLATED) as bundle:
            bundle.writestr("README.txt", self._report_export_readme(disease))
            bundle.writestr("report_metadata.json", self._json_text(metadata))
            bundle.writestr(f"report/physician_compare_{safe_disease}.docx", report_docx)
            for name, payload in figures.items():
                bundle.writestr(f"figures/{name}", payload)
            for name, rows in table_payload.items():
                fieldnames = list(rows[0].keys()) if rows else ["note"]
                normalized_rows = rows or [{"note": "暂无数据"}]
                bundle.writestr(f"tables/{name}", self._csv_text(normalized_rows, fieldnames))

        filename = f"physician_compare_report_{safe_disease}.zip"
        return filename, archive.getvalue(), final_timings

    def _build_doctor_subgraphs(
        self,
        disease: str,
        contexts: list[dict[str, object]],
    ) -> dict[str, dict[str, object]]:
        # 按医家逐个构建子图。
        # 上游 contexts 已经从 Neo4j 准备好了当前医家、当前病名、病例范围和图快照。
        grouped: dict[str, dict[str, object]] = {}
        for context in contexts:
            doctor = context.get("doctor")
            disease_node = context.get("disease")
            snapshot = context.get("snapshot")
            source_cases = context.get("source_cases")
            if not isinstance(doctor, GraphNode) or not isinstance(snapshot, GraphSnapshot):
                continue
            typed_disease = disease_node if isinstance(disease_node, GraphNode) else self._disease_node(snapshot, disease)
            if typed_disease is None:
                continue
            case_list = source_cases if isinstance(source_cases, list) else []
            grouped[doctor.name] = self._extract_core_subgraph(doctor, typed_disease, snapshot, case_list)
        return grouped

    def _extract_core_subgraph(
        self,
        doctor: GraphNode,
        disease: GraphNode,
        snapshot: GraphSnapshot,
        source_cases: list[str],
    ) -> dict[str, object]:
        # 核心子图抽取规则：
        # 1. 当前医家是一个锚点
        # 2. 当前病名是一个锚点
        # 3. 只有同时被“当前医家”和“当前病名”指向的 C/D/E 节点，才能进核心子图
        # 4. 只有这些已纳入节点之间真实存在的结构边，才能保留
        node_by_id = {node.id: node for node in snapshot.nodes}
        filtered_edges = [
            edge
            for edge in snapshot.edges
            if not source_cases or any(source_case in edge.source_cases for source_case in source_cases)
        ]
        # doctor_nodes / disease_nodes 先分别收集两侧证据，后面再取交集。
        doctor_nodes: dict[str, dict[tuple[str, str], GraphNode]] = {category: {} for category in self.SUBGRAPH_CATEGORY_ORDER}
        disease_nodes: dict[str, dict[tuple[str, str], GraphNode]] = {category: {} for category in self.SUBGRAPH_CATEGORY_ORDER}
        doctor_evidence: dict[tuple[str, str], set[str]] = {}
        disease_evidence: dict[tuple[str, str], set[str]] = {}
        doctor_disease_edge_present = False

        for edge in filtered_edges:
            source_node = node_by_id.get(edge.source)
            target_node = node_by_id.get(edge.target)
            if source_node is None or target_node is None:
                continue
            if {source_node.id, target_node.id} == {doctor.id, disease.id}:
                doctor_disease_edge_present = True
            if doctor.id in {source_node.id, target_node.id}:
                feature_node = target_node if source_node.id == doctor.id else source_node
                if feature_node.type in self.SUBGRAPH_FEATURE_TYPES:
                    category = self.RESULT_TYPES[feature_node.type]
                    self._store_canonical_node(doctor_nodes[category], feature_node)
                    doctor_evidence.setdefault(self._canonical_key(feature_node), set()).add(edge.type)
            if disease.id in {source_node.id, target_node.id}:
                feature_node = target_node if source_node.id == disease.id else source_node
                if feature_node.type in self.SUBGRAPH_FEATURE_TYPES:
                    category = self.RESULT_TYPES[feature_node.type]
                    self._store_canonical_node(disease_nodes[category], feature_node)
                    disease_evidence.setdefault(self._canonical_key(feature_node), set()).add(edge.type)

        # 第 1 步：只保留“医家边”和“病名边”同时成立的特征节点。
        node_map: dict[str, dict[tuple[str, str], GraphNode]] = {category: {} for category in self.SUBGRAPH_CATEGORY_ORDER}
        audit_nodes: list[PhysicianSubgraphAuditNode] = []
        included_feature_keys: set[tuple[str, str]] = set()
        for category in self.SUBGRAPH_CATEGORY_ORDER:
            shared_keys = set(doctor_nodes[category]) & set(disease_nodes[category])
            for key in sorted(shared_keys, key=lambda item: item[1]):
                left_node = doctor_nodes[category][key]
                right_node = disease_nodes[category][key]
                chosen = left_node if left_node.id <= right_node.id else right_node
                node_map[category][key] = chosen
                included_feature_keys.add(key)
                doctor_relations = "、".join(sorted(doctor_evidence.get(key, set())))
                disease_relations = "、".join(sorted(disease_evidence.get(key, set())))
                audit_nodes.append(
                    PhysicianSubgraphAuditNode(
                        id=chosen.id,
                        name=chosen.name,
                        label=chosen.label,
                        type=chosen.type,
                        inclusion_reason=f"同时满足 {doctor.name} 的 {doctor_relations} 与 {disease.name} 的 {disease_relations}",
                    )
                )

        # 第 2 步：生成核心子图里的边。
        # concrete_edge_map 用真实医家名，给页面展示和导出使用。
        # shared_edge_map 把医家名规整成“当前医家”，方便做多医家共同边比较。
        concrete_edge_map: dict[str, PhysicianSubgraphEdge] = {}
        shared_edge_map: dict[str, PhysicianSubgraphEdge] = {}
        audit_edges: list[PhysicianSubgraphAuditEdge] = []

        if doctor_disease_edge_present:
            signature = f"A医家-B病名|{disease.name}"
            concrete_edge_map[signature] = self._make_subgraph_edge(
                signature=signature,
                relation_type="A医家-B病名",
                source_name=doctor.name,
                source_type=doctor.type,
                target_name=disease.name,
                target_type=disease.type,
            )
            shared_edge_map[signature] = self._make_subgraph_edge(
                signature=signature,
                relation_type="A医家-B病名",
                source_name="当前医家",
                source_type=doctor.type,
                target_name=disease.name,
                target_type=disease.type,
            )
            audit_edges.append(self._make_subgraph_audit_edge(concrete_edge_map[signature]))

        for category in self.SUBGRAPH_CATEGORY_ORDER:
            for key, feature_node in sorted(node_map[category].items(), key=lambda item: item[1].name):
                for relation_type in sorted(doctor_evidence.get(key, set())):
                    signature = f"{relation_type}|当前医家|{feature_node.type}|{feature_node.name}"
                    concrete_edge_map[signature] = self._make_subgraph_edge(
                        signature=signature,
                        relation_type=relation_type,
                        source_name=doctor.name,
                        source_type=doctor.type,
                        target_name=feature_node.name,
                        target_type=feature_node.type,
                    )
                    shared_edge_map[signature] = self._make_subgraph_edge(
                        signature=signature,
                        relation_type=relation_type,
                        source_name="当前医家",
                        source_type=doctor.type,
                        target_name=feature_node.name,
                        target_type=feature_node.type,
                    )
                    audit_edges.append(self._make_subgraph_audit_edge(concrete_edge_map[signature]))
                for relation_type in sorted(disease_evidence.get(key, set())):
                    signature = f"{relation_type}|{disease.name}|{feature_node.type}|{feature_node.name}"
                    concrete_edge_map[signature] = self._make_subgraph_edge(
                        signature=signature,
                        relation_type=relation_type,
                        source_name=disease.name,
                        source_type=disease.type,
                        target_name=feature_node.name,
                        target_type=feature_node.type,
                    )
                    shared_edge_map[signature] = concrete_edge_map[signature]
                    audit_edges.append(self._make_subgraph_audit_edge(concrete_edge_map[signature]))

        # 第 3 步：补上已纳入核心节点之间真实存在的结构边。
        for edge in filtered_edges:
            source_node = node_by_id.get(edge.source)
            target_node = node_by_id.get(edge.target)
            if source_node is None or target_node is None:
                continue
            source_key = self._canonical_key(source_node)
            target_key = self._canonical_key(target_node)
            if source_key not in included_feature_keys or target_key not in included_feature_keys:
                continue
            if source_node.type not in self.SUBGRAPH_FEATURE_TYPES or target_node.type not in self.SUBGRAPH_FEATURE_TYPES:
                continue
            if edge.type not in self.PATH_RELATION_TYPES:
                continue
            ordered_source, ordered_target = self._ordered_relation_nodes(source_node, target_node, edge.type)
            signature = self._subgraph_structure_signature(edge.type, ordered_source, ordered_target)
            concrete_edge = self._make_subgraph_edge(
                signature=signature,
                relation_type=edge.type,
                source_name=ordered_source.name,
                source_type=ordered_source.type,
                target_name=ordered_target.name,
                target_type=ordered_target.type,
            )
            concrete_edge_map[signature] = concrete_edge
            shared_edge_map[signature] = concrete_edge
            audit_edges.append(self._make_subgraph_audit_edge(concrete_edge))

        # 第 4 步：统计关系分布，并把核心子图转换成向量算法要用的 NetworkX 图。
        relation_distribution = self._relation_distribution(list(concrete_edge_map.values()))
        audit_edge_map = {
            (item.relation_type, item.text): item
            for item in audit_edges
        }
        audit_edges = sorted(audit_edge_map.values(), key=lambda item: (item.relation_type, item.text))
        return {
            "doctor": doctor,
            "disease": disease,
            "node_map": node_map,
            "node_keys": {key for category in self.SUBGRAPH_CATEGORY_ORDER for key in node_map[category]},
            "edge_signatures": set(shared_edge_map),
            "shared_edge_map": shared_edge_map,
            "concrete_edge_map": concrete_edge_map,
            "vector_graph": self._build_relation_node_vector_graph(concrete_edge_map),
            "relation_distribution": relation_distribution,
            "audit_nodes": audit_nodes,
            "audit_edges": audit_edges,
        }

    def _collect_doctor_node_maps(self, contexts: list[dict[str, object]]) -> dict[str, dict[str, dict[tuple[str, str], GraphNode]]]:
        doctor_node_maps: dict[str, dict[str, dict[tuple[str, str], GraphNode]]] = {}
        for context in contexts:
            doctor = context.get("doctor")
            snapshot = context.get("snapshot")
            source_cases = context.get("source_cases")
            if not isinstance(doctor, GraphNode) or not isinstance(snapshot, GraphSnapshot):
                continue
            case_list = source_cases if isinstance(source_cases, list) else []
            doctor_node_maps[doctor.name] = self._doctor_category_node_map(doctor, snapshot, case_list)
        return doctor_node_maps

    def _doctor_category_node_map(
        self,
        doctor: GraphNode,
        snapshot: GraphSnapshot,
        source_cases: list[str],
    ) -> dict[str, dict[tuple[str, str], GraphNode]]:
        node_by_id = {node.id: node for node in snapshot.nodes}
        grouped: dict[str, dict[tuple[str, str], GraphNode]] = {
            "patterns": {},
            "causes": {},
            "mechanisms": {},
        }
        for edge in snapshot.edges:
            if source_cases and not any(source_case in edge.source_cases for source_case in source_cases):
                continue
            feature_id = ""
            if edge.source == doctor.id:
                feature_id = edge.target
            elif edge.target == doctor.id:
                feature_id = edge.source
            if not feature_id:
                continue
            node = node_by_id.get(feature_id)
            if node is None or node.type not in self.RESULT_TYPES:
                continue
            category = self.RESULT_TYPES[node.type]
            key = self._canonical_key(node)
            existing = grouped[category].get(key)
            if existing is None or node.id < existing.id:
                grouped[category][key] = node
        return grouped

    def _build_shared_nodes(
        self,
        doctor_node_maps: dict[str, dict[str, dict[tuple[str, str], GraphNode]]],
    ) -> SharedCompareNodeGroup:
        grouped: dict[str, list[SharedCompareNode]] = {
            "patterns": [],
            "causes": [],
            "mechanisms": [],
        }
        for category in grouped:
            doctor_by_key: dict[tuple[str, str], list[str]] = {}
            node_by_key: dict[tuple[str, str], GraphNode] = {}
            for doctor_name, category_map in doctor_node_maps.items():
                node_map = category_map[category]
                for key, node in node_map.items():
                    doctor_by_key.setdefault(key, []).append(doctor_name)
                    existing = node_by_key.get(key)
                    if existing is None or node.id < existing.id:
                        node_by_key[key] = node
            grouped[category] = [
                SharedCompareNode(
                    node=self._compare_node(node_by_key[key]),
                    doctors=sorted(doctors),
                )
                for key, doctors in doctor_by_key.items()
                if len(doctors) >= 2
            ]
            grouped[category].sort(key=lambda item: (-len(item.doctors), item.node.name))
        return SharedCompareNodeGroup(**grouped)

    def _build_doctor_profiles(
        self,
        contexts: list[dict[str, object]],
        doctor_node_maps: dict[str, dict[str, dict[tuple[str, str], GraphNode]]],
        shared_nodes: SharedCompareNodeGroup,
    ) -> list[PhysicianNodeProfile]:
        shared_keys = {
            category: {
                self._canonical_key_from_compare_node(item.node)
                for item in getattr(shared_nodes, category)
            }
            for category in ("patterns", "causes", "mechanisms")
        }
        profiles: list[PhysicianNodeProfile] = []
        for context in contexts:
            doctor = context.get("doctor")
            if not isinstance(doctor, GraphNode):
                continue
            category_map = doctor_node_maps.get(doctor.name, {"patterns": {}, "causes": {}, "mechanisms": {}})
            all_nodes = self._compare_node_group(category_map)
            shared_group = self._compare_node_group(
                {
                    category: {
                        key: node
                        for key, node in category_map[category].items()
                        if key in shared_keys[category]
                    }
                    for category in ("patterns", "causes", "mechanisms")
                }
            )
            unique_group = self._compare_node_group(
                {
                    category: {
                        key: node
                        for key, node in category_map[category].items()
                        if key not in shared_keys[category]
                    }
                    for category in ("patterns", "causes", "mechanisms")
                }
            )
            profiles.append(
                PhysicianNodeProfile(
                    doctor=doctor,
                    all=all_nodes,
                    shared=shared_group,
                    unique=unique_group,
                )
            )
        return profiles

    def _build_similarity_group(
        self,
        doctors: list[GraphNode],
        similarity_rows: dict[str, list[dict[str, object]]],
    ) -> PhysicianSimilarityGroup:
        doctor_names = sorted({doctor.name for doctor in doctors})
        grouped = {}
        for category in ("patterns", "causes", "mechanisms", "overall"):
            grouped[category] = self._complete_similarity_pairs(
                doctor_names,
                similarity_rows.get(category, []),
            )
        return PhysicianSimilarityGroup(**grouped)

    def _complete_similarity_pairs(
        self,
        doctor_names: list[str],
        rows: list[dict[str, object]],
    ) -> list[PhysicianSimilarityPair]:
        row_map = {
            tuple(sorted((str(row.get("left_doctor", "")), str(row.get("right_doctor", ""))))): row
            for row in rows
            if row.get("left_doctor") and row.get("right_doctor")
        }
        pairs: list[PhysicianSimilarityPair] = []
        for left_index, left_doctor in enumerate(doctor_names):
            for right_doctor in doctor_names[left_index + 1:]:
                row = row_map.get((left_doctor, right_doctor), {})
                pairs.append(
                    PhysicianSimilarityPair(
                        left_doctor=left_doctor,
                        right_doctor=right_doctor,
                        jaccard=round(float(row.get("jaccard", 0.0)), 6),
                        overlap=round(float(row.get("overlap", 0.0)), 6),
                        cosine=round(float(row.get("cosine", 0.0)), 6),
                    )
                )
        return pairs

    def _build_rwr_results(self, disease: str, contexts: list[dict[str, object]]) -> list[PhysicianNodeRwrResult]:
        results: list[PhysicianNodeRwrResult] = []
        for context in contexts:
            doctor = context.get("doctor")
            snapshot = context.get("snapshot")
            if not isinstance(doctor, GraphNode) or not isinstance(snapshot, GraphSnapshot):
                continue
            disease_node = self._disease_node(snapshot, disease)
            if disease_node is None:
                continue
            graph = self._build_graph(snapshot)
            if doctor.id not in graph or disease_node.id not in graph:
                continue
            rankings = self._canonical_rwr_rankings(
                graph,
                doctor.id,
                disease_node.id,
                self.PRIMARY_RESTART_PROBABILITY,
            )
            results.append(
                PhysicianNodeRwrResult(
                    doctor=doctor,
                    restart_probability=self.PRIMARY_RESTART_PROBABILITY,
                    rankings=rankings,
                )
            )
        return results

    def _build_doctor_paths(self, contexts: list[dict[str, object]]) -> dict[str, dict[str, list[PhysicianPathChain] | set[str]]]:
        grouped: dict[str, dict[str, list[PhysicianPathChain] | set[str]]] = {}
        for context in contexts:
            doctor = context.get("doctor")
            snapshot = context.get("snapshot")
            source_cases = context.get("source_cases")
            if not isinstance(doctor, GraphNode) or not isinstance(snapshot, GraphSnapshot):
                continue
            case_list = source_cases if isinstance(source_cases, list) else []
            grouped[doctor.name] = self._extract_doctor_paths(doctor, snapshot, case_list)
        return grouped

    def _extract_doctor_paths(
        self,
        doctor: GraphNode,
        snapshot: GraphSnapshot,
        source_cases: list[str],
    ) -> dict[str, list[PhysicianPathChain] | set[str]]:
        node_by_id = {node.id: node for node in snapshot.nodes}
        direct_causes: dict[tuple[str, str], GraphNode] = {}
        direct_mechanisms: dict[tuple[str, str], GraphNode] = {}
        direct_patterns: dict[tuple[str, str], GraphNode] = {}
        de_pairs: set[tuple[tuple[str, str], tuple[str, str]]] = set()
        ce_pairs: set[tuple[tuple[str, str], tuple[str, str]]] = set()
        cd_pairs: set[tuple[tuple[str, str], tuple[str, str]]] = set()

        for edge in snapshot.edges:
            if source_cases and not any(source_case in edge.source_cases for source_case in source_cases):
                continue
            source_node = node_by_id.get(edge.source)
            target_node = node_by_id.get(edge.target)
            if source_node is None or target_node is None:
                continue
            if doctor.id in {edge.source, edge.target}:
                feature_node = target_node if edge.source == doctor.id else source_node
                if feature_node.type not in self.PATH_DIRECT_TYPES:
                    continue
                target_map = self._path_node_map_for_type(
                    feature_node.type,
                    direct_patterns,
                    direct_causes,
                    direct_mechanisms,
                )
                self._store_canonical_node(target_map, feature_node)
                continue

            node_types = {source_node.type, target_node.type}
            if node_types == {"D病因", "E病机"}:
                cause_node = source_node if source_node.type == "D病因" else target_node
                mechanism_node = target_node if target_node.type == "E病机" else source_node
                de_pairs.add((self._canonical_key(cause_node), self._canonical_key(mechanism_node)))
                continue
            if node_types == {"C证型", "E病机"}:
                pattern_node = source_node if source_node.type == "C证型" else target_node
                mechanism_node = target_node if target_node.type == "E病机" else source_node
                ce_pairs.add((self._canonical_key(pattern_node), self._canonical_key(mechanism_node)))
                continue
            if node_types == {"C证型", "D病因"}:
                pattern_node = source_node if source_node.type == "C证型" else target_node
                cause_node = target_node if target_node.type == "D病因" else source_node
                cd_pairs.add((self._canonical_key(cause_node), self._canonical_key(pattern_node)))

        complete_paths: list[PhysicianPathChain] = []
        used_de_pairs: set[tuple[tuple[str, str], tuple[str, str]]] = set()
        used_ce_pairs: set[tuple[tuple[str, str], tuple[str, str]]] = set()
        used_cause_keys: set[tuple[str, str]] = set()
        used_mechanism_keys: set[tuple[str, str]] = set()
        used_pattern_keys: set[tuple[str, str]] = set()

        for cause_key, cause_node in direct_causes.items():
            for mechanism_key, mechanism_node in direct_mechanisms.items():
                if (cause_key, mechanism_key) not in de_pairs:
                    continue
                for pattern_key, pattern_node in direct_patterns.items():
                    if (pattern_key, mechanism_key) not in ce_pairs:
                        continue
                    complete_paths.append(
                        self._make_path_chain(
                            path_type="D-E-C",
                            path_category="complete",
                            cause=cause_node,
                            mechanism=mechanism_node,
                            pattern=pattern_node,
                        )
                    )
                    used_de_pairs.add((cause_key, mechanism_key))
                    used_ce_pairs.add((pattern_key, mechanism_key))
                    used_cause_keys.add(cause_key)
                    used_mechanism_keys.add(mechanism_key)
                    used_pattern_keys.add(pattern_key)

        partial_paths: list[PhysicianPathChain] = []
        for cause_key, mechanism_key in sorted(de_pairs):
            if cause_key not in direct_causes or mechanism_key not in direct_mechanisms:
                continue
            if (cause_key, mechanism_key) in used_de_pairs:
                continue
            partial_paths.append(
                self._make_path_chain(
                    path_type="D-E",
                    path_category="partial",
                    cause=direct_causes[cause_key],
                    mechanism=direct_mechanisms[mechanism_key],
                )
            )
            used_cause_keys.add(cause_key)
            used_mechanism_keys.add(mechanism_key)

        for pattern_key, mechanism_key in sorted(ce_pairs):
            if pattern_key not in direct_patterns or mechanism_key not in direct_mechanisms:
                continue
            if (pattern_key, mechanism_key) in used_ce_pairs:
                continue
            partial_paths.append(
                self._make_path_chain(
                    path_type="E-C",
                    path_category="partial",
                    mechanism=direct_mechanisms[mechanism_key],
                    pattern=direct_patterns[pattern_key],
                )
            )
            used_pattern_keys.add(pattern_key)
            used_mechanism_keys.add(mechanism_key)

        for cause_key, pattern_key in sorted(cd_pairs):
            if cause_key not in direct_causes or pattern_key not in direct_patterns:
                continue
            partial_paths.append(
                self._make_path_chain(
                    path_type="D-C",
                    path_category="partial",
                    cause=direct_causes[cause_key],
                    pattern=direct_patterns[pattern_key],
                )
            )
            used_cause_keys.add(cause_key)
            used_pattern_keys.add(pattern_key)

        single_paths: list[PhysicianPathChain] = []
        for cause_key, cause_node in sorted(direct_causes.items(), key=lambda item: item[1].name):
            if cause_key in used_cause_keys:
                continue
            single_paths.append(self._make_path_chain(path_type="D", path_category="single", cause=cause_node))
        for mechanism_key, mechanism_node in sorted(direct_mechanisms.items(), key=lambda item: item[1].name):
            if mechanism_key in used_mechanism_keys:
                continue
            single_paths.append(self._make_path_chain(path_type="E", path_category="single", mechanism=mechanism_node))
        for pattern_key, pattern_node in sorted(direct_patterns.items(), key=lambda item: item[1].name):
            if pattern_key in used_pattern_keys:
                continue
            single_paths.append(self._make_path_chain(path_type="C", path_category="single", pattern=pattern_node))

        all_signatures = {path.signature for path in complete_paths + partial_paths + single_paths}
        total_direct_node_count = len(direct_causes) + len(direct_mechanisms) + len(direct_patterns)
        used_direct_node_count = len(used_cause_keys | used_mechanism_keys | used_pattern_keys)
        return {
            "complete_paths": complete_paths,
            "partial_paths": partial_paths,
            "single_paths": single_paths,
            "all_signatures": all_signatures,
            "total_direct_node_count": total_direct_node_count,
            "used_direct_node_count": used_direct_node_count,
        }

    def _build_shared_paths(
        self,
        doctor_paths: dict[str, dict[str, list[PhysicianPathChain] | set[str]]],
    ) -> list[SharedPhysicianPath]:
        path_by_signature: dict[str, PhysicianPathChain] = {}
        doctors_by_signature: dict[str, list[str]] = {}
        for doctor_name, payload in doctor_paths.items():
            for bucket in ("complete_paths", "partial_paths", "single_paths"):
                for path in payload[bucket]:
                    path_by_signature[path.signature] = path
                    doctors_by_signature.setdefault(path.signature, []).append(doctor_name)
        shared_paths = [
            SharedPhysicianPath(
                path=path_by_signature[signature],
                doctors=sorted(doctors),
            )
            for signature, doctors in doctors_by_signature.items()
            if len(doctors) >= 2
        ]
        shared_paths.sort(key=lambda item: (-len(item.doctors), item.path.path_type, item.path.text))
        return shared_paths

    def _build_path_profiles(
        self,
        contexts: list[dict[str, object]],
        doctor_paths: dict[str, dict[str, list[PhysicianPathChain] | set[str]]],
        shared_paths: list[SharedPhysicianPath],
    ) -> list[PhysicianPathProfile]:
        shared_signatures = {item.path.signature for item in shared_paths}
        profiles: list[PhysicianPathProfile] = []
        for context in contexts:
            doctor = context.get("doctor")
            if not isinstance(doctor, GraphNode):
                continue
            payload = doctor_paths.get(
                doctor.name,
                {"complete_paths": [], "partial_paths": [], "single_paths": [], "all_signatures": set()},
            )
            complete_paths = list(payload["complete_paths"])
            partial_paths = list(payload["partial_paths"])
            single_paths = list(payload["single_paths"])
            unique_paths = [
                path
                for path in complete_paths + partial_paths + single_paths
                if path.signature not in shared_signatures
            ]
            total_count = len(complete_paths) + len(partial_paths) + len(single_paths)
            profiles.append(
                PhysicianPathProfile(
                    doctor=doctor,
                    complete_paths=complete_paths,
                    partial_paths=partial_paths,
                    single_paths=single_paths,
                    unique_paths=unique_paths,
                    completeness=PhysicianPathCompleteness(
                        complete_count=len(complete_paths),
                        partial_count=len(partial_paths),
                        single_count=len(single_paths),
                        total_count=total_count,
                        complete_ratio=round(len(complete_paths) / total_count, 4) if total_count else 0.0,
                        path_coverage=round(
                            int(payload.get("used_direct_node_count", 0)) / int(payload.get("total_direct_node_count", 0)),
                            4,
                        ) if int(payload.get("total_direct_node_count", 0)) else 0.0,
                    ),
                )
            )
        return profiles

    def _build_path_embedding_profiles(
        self,
        doctors: list[GraphNode],
        doctor_paths: dict[str, dict[str, list[PhysicianPathChain] | set[str]]],
    ) -> list[PhysicianPathEmbeddingProfile]:
        Word2Vec = self._load_word2vec_class()

        walks: list[list[str]] = []
        walks_by_doctor: dict[str, list[list[str]]] = {}
        doctor_names = sorted({doctor.name for doctor in doctors})
        for doctor_name in doctor_names:
            payload = doctor_paths.get(doctor_name, {})
            doctor_walks: list[list[str]] = []
            for bucket in ("complete_paths", "partial_paths", "single_paths"):
                for path in payload.get(bucket, []):
                    if not isinstance(path, PhysicianPathChain):
                        continue
                    walk = self._metapath_walk(path)
                    if not walk:
                        continue
                    walks.append(walk)
                    doctor_walks.append(walk)
            walks_by_doctor[doctor_name] = doctor_walks

        if not walks:
            return [PhysicianPathEmbeddingProfile(doctor=doctor, vector=[]) for doctor in doctors]

        model = Word2Vec(
            sentences=walks,
            vector_size=64,
            window=3,
            min_count=1,
            sg=1,
            workers=1,
            epochs=10,
            seed=self.RANDOM_SEED,
        )
        doctor_by_name = {doctor.name: doctor for doctor in doctors}
        profiles: list[PhysicianPathEmbeddingProfile] = []
        for doctor_name in doctor_names:
            path_vectors: list[np.ndarray] = []
            for walk in walks_by_doctor.get(doctor_name, []):
                token_vectors = [model.wv[token] for token in walk if token in model.wv]
                if token_vectors:
                    path_vectors.append(np.mean(np.asarray(token_vectors, dtype=float), axis=0))
            vector = np.mean(np.asarray(path_vectors, dtype=float), axis=0) if path_vectors else np.zeros(model.vector_size)
            profiles.append(
                PhysicianPathEmbeddingProfile(
                    doctor=doctor_by_name[doctor_name],
                    vector=self._float_list(vector),
                )
            )
        return profiles

    def _build_path_similarity_pairs(
        self,
        doctors: list[GraphNode],
        doctor_paths: dict[str, dict[str, list[PhysicianPathChain] | set[str]]],
        embeddings: list[PhysicianPathEmbeddingProfile],
    ) -> list[PhysicianPathSimilarityPair]:
        pairs: list[PhysicianPathSimilarityPair] = []
        doctor_names = sorted({doctor.name for doctor in doctors})
        embedding_map = {
            item.doctor.name: np.asarray(item.vector, dtype=float)
            for item in embeddings
            if item.vector
        }
        for left_index, left_doctor in enumerate(doctor_names):
            left_signatures = set(doctor_paths.get(left_doctor, {}).get("all_signatures", set()))
            for right_doctor in doctor_names[left_index + 1:]:
                right_signatures = set(doctor_paths.get(right_doctor, {}).get("all_signatures", set()))
                union = left_signatures | right_signatures
                shared = left_signatures & right_signatures
                pairs.append(
                    PhysicianPathSimilarityPair(
                        left_doctor=left_doctor,
                        right_doctor=right_doctor,
                        shared_path_count=len(shared),
                        union_path_count=len(union),
                        path_jaccard=round(len(shared) / len(union), 6) if union else 0.0,
                        metapath2vec_cosine=self._cosine_similarity(
                            embedding_map.get(left_doctor, np.asarray([], dtype=float)),
                            embedding_map.get(right_doctor, np.asarray([], dtype=float)),
                        ),
                    )
                )
        return pairs

    def _build_shared_subgraph_edges(
        self,
        subgraphs: dict[str, dict[str, object]],
    ) -> list[SharedPhysicianSubgraphEdge]:
        edge_by_signature: dict[str, PhysicianSubgraphEdge] = {}
        doctors_by_signature: dict[str, list[str]] = {}
        for doctor_name, payload in subgraphs.items():
            shared_edge_map = payload.get("shared_edge_map", {})
            if not isinstance(shared_edge_map, dict):
                continue
            for signature, edge in shared_edge_map.items():
                if not isinstance(edge, PhysicianSubgraphEdge):
                    continue
                edge_by_signature[signature] = edge
                doctors_by_signature.setdefault(signature, []).append(doctor_name)
        shared_edges = [
            SharedPhysicianSubgraphEdge(
                edge=edge_by_signature[signature],
                doctors=sorted(doctors),
            )
            for signature, doctors in doctors_by_signature.items()
            if len(doctors) >= 2
        ]
        shared_edges.sort(key=lambda item: (-len(item.doctors), item.edge.relation_type, item.edge.text))
        return shared_edges

    def _build_subgraph_profiles(
        self,
        contexts: list[dict[str, object]],
        subgraphs: dict[str, dict[str, object]],
        shared_nodes: SharedCompareNodeGroup,
        shared_edges: list[SharedPhysicianSubgraphEdge],
    ) -> list[PhysicianSubgraphProfile]:
        shared_node_keys = {
            category: {self._canonical_key_from_compare_node(item.node) for item in getattr(shared_nodes, category)}
            for category in self.SUBGRAPH_CATEGORY_ORDER
        }
        shared_edge_signatures = {item.edge.signature for item in shared_edges}
        profiles: list[PhysicianSubgraphProfile] = []
        for context in contexts:
            doctor = context.get("doctor")
            if not isinstance(doctor, GraphNode):
                continue
            payload = subgraphs.get(doctor.name)
            if not isinstance(payload, dict):
                continue
            node_map = payload.get("node_map", {})
            concrete_edge_map = payload.get("concrete_edge_map", {})
            audit_nodes = payload.get("audit_nodes", [])
            audit_edges = payload.get("audit_edges", [])
            relation_distribution = payload.get("relation_distribution", [])
            if not isinstance(node_map, dict) or not isinstance(concrete_edge_map, dict):
                continue
            unique_node_map = {
                category: {
                    key: node
                    for key, node in node_map.get(category, {}).items()
                    if key not in shared_node_keys[category]
                }
                for category in self.SUBGRAPH_CATEGORY_ORDER
            }
            concrete_edges = [
                edge
                for edge in sorted(concrete_edge_map.values(), key=lambda item: (item.relation_type, item.text))
                if isinstance(edge, PhysicianSubgraphEdge)
            ]
            unique_edges = [edge for edge in concrete_edges if edge.signature not in shared_edge_signatures]
            profiles.append(
                PhysicianSubgraphProfile(
                    doctor=doctor,
                    node_count=sum(len(node_map.get(category, {})) for category in self.SUBGRAPH_CATEGORY_ORDER),
                    edge_count=len(concrete_edges),
                    nodes=self._compare_node_group(node_map),
                    unique_nodes=self._compare_node_group(unique_node_map),
                    edges=concrete_edges,
                    unique_edges=unique_edges,
                    relation_distribution=[
                        item for item in relation_distribution if isinstance(item, PhysicianSubgraphRelationStat)
                    ],
                    audit_nodes=[
                        item
                        for item in sorted(
                            [node for node in audit_nodes if isinstance(node, PhysicianSubgraphAuditNode)],
                            key=lambda node: (node.type, node.name),
                        )
                        if isinstance(item, PhysicianSubgraphAuditNode)
                    ],
                    audit_edges=[
                        item for item in audit_edges
                        if isinstance(item, PhysicianSubgraphAuditEdge)
                    ],
                )
            )
        return profiles

    def _build_subgraph_similarity_pairs(
        self,
        doctors: list[GraphNode],
        subgraphs: dict[str, dict[str, object]],
        embedding_payload: dict[str, object],
    ) -> list[PhysicianSubgraphSimilarityPair]:
        pairs: list[PhysicianSubgraphSimilarityPair] = []
        doctor_names = sorted({doctor.name for doctor in doctors})
        vector_embeddings = self._typed_mapping(embedding_payload.get("vector_embeddings"))
        for left_index, left_doctor in enumerate(doctor_names):
            left_payload = subgraphs.get(left_doctor, {})
            left_nodes = set(left_payload.get("node_keys", set()))
            left_edges = set(left_payload.get("edge_signatures", set()))
            for right_doctor in doctor_names[left_index + 1:]:
                right_payload = subgraphs.get(right_doctor, {})
                right_nodes = set(right_payload.get("node_keys", set()))
                right_edges = set(right_payload.get("edge_signatures", set()))
                node_jaccard = self._jaccard(left_nodes, right_nodes)
                edge_jaccard = self._jaccard(left_edges, right_edges)
                combined_left = {f"node:{item}" for item in left_nodes} | {f"edge:{item}" for item in left_edges}
                combined_right = {f"node:{item}" for item in right_nodes} | {f"edge:{item}" for item in right_edges}
                pairs.append(
                    PhysicianSubgraphSimilarityPair(
                        left_doctor=left_doctor,
                        right_doctor=right_doctor,
                        node_jaccard=node_jaccard,
                        edge_jaccard=edge_jaccard,
                        subgraph_jaccard=self._jaccard(combined_left, combined_right),
                        graph2vec_cosine=self._vector_similarity(vector_embeddings, "Graph2Vec", left_doctor, right_doctor),
                    )
                )
        return pairs

    def _build_relation_node_vector_graph(
        self,
        concrete_edge_map: dict[str, PhysicianSubgraphEdge],
    ) -> nx.Graph:
        # 这一步把核心子图变成普通 NetworkX 图，给 Graph2Vec 使用。
        # 关键不是“原样复制”，而是“关系节点化”：
        #   实体节点 -- 关系节点 -- 实体节点
        # 这样能尽量保留关系类型在结构中的位置。
        graph = nx.Graph()
        for edge in concrete_edge_map.values():
            if not isinstance(edge, PhysicianSubgraphEdge):
                continue
            source_key = f"entity|{edge.source_type}|{edge.source_name}"
            target_key = f"entity|{edge.target_type}|{edge.target_name}"
            relation_key = (
                f"relation|{edge.relation_type}|{edge.source_type}|{edge.source_name}|"
                f"{edge.target_type}|{edge.target_name}"
            )
            graph.add_node(source_key, kind="entity", entity_type=edge.source_type, name=edge.source_name)
            graph.add_node(target_key, kind="entity", entity_type=edge.target_type, name=edge.target_name)
            graph.add_node(
                relation_key,
                kind="relation",
                relation_type=edge.relation_type,
                source_name=edge.source_name,
                source_type=edge.source_type,
                target_name=edge.target_name,
                target_type=edge.target_type,
            )
            graph.add_edge(source_key, relation_key)
            graph.add_edge(relation_key, target_key)
        # 最后统一改成整数节点 id，避免图嵌入库处理字符串节点时出现兼容问题。
        return nx.convert_node_labels_to_integers(graph, ordering="sorted", label_attribute="original_key")

    def _build_subgraph_embeddings(
        self,
        doctor_names: list[str],
        subgraphs: dict[str, dict[str, object]],
    ) -> dict[str, object]:
        try:
            from karateclub import Graph2Vec
        except ImportError:
            try:
                from karateclub.graph_embedding import Graph2Vec
            except ModuleNotFoundError as exc:
                raise ValueError("子图向量比较依赖 karateclub，请先进入 KG 环境再运行。") from exc
        except ModuleNotFoundError as exc:
            raise ValueError("子图向量比较依赖 karateclub，请先进入 KG 环境再运行。") from exc

        self._ensure_karateclub_runtime()

        graphs: list[nx.Graph] = []
        valid_doctors: list[str] = []
        for doctor_name in doctor_names:
            payload = subgraphs.get(doctor_name, {})
            vector_graph = payload.get("vector_graph")
            if isinstance(vector_graph, nx.Graph):
                graphs.append(vector_graph)
                valid_doctors.append(doctor_name)
        if not graphs:
            return {"vector_embeddings": {"Graph2Vec": {}}}

        embeddings: dict[str, dict[str, np.ndarray]] = {}
        graph2vec_model = Graph2Vec(
            min_count=1,
            workers=1,
            seed=self.RANDOM_SEED,
        )
        graph2vec_model.fit(graphs)
        embeddings["Graph2Vec"] = self._named_embeddings(valid_doctors, graph2vec_model.get_embedding())
        return {
            "vector_embeddings": embeddings,
        }

    @staticmethod
    def _ensure_karateclub_runtime() -> None:
        try:
            import scipy as sp
        except ModuleNotFoundError as exc:
            raise ValueError("子图向量比较依赖 scipy，请先在 KG 环境补齐依赖。") from exc
        if not hasattr(sp, "errstate"):
            sp.errstate = np.errstate

    @staticmethod
    def _named_embeddings(doctor_names: list[str], matrix: object) -> dict[str, np.ndarray]:
        array = np.asarray(matrix, dtype=float)
        if array.ndim == 1:
            array = np.expand_dims(array, axis=0)
        return {
            doctor_name: array[index]
            for index, doctor_name in enumerate(doctor_names)
            if index < len(array)
        }

    def _vector_similarity(
        self,
        embeddings: dict[str, dict[str, np.ndarray]],
        algorithm: str,
        left_doctor: str,
        right_doctor: str,
    ) -> float:
        # 先取出两个医家的向量，再算余弦相似度。
        algorithm_embeddings = embeddings.get(algorithm, {})
        left = algorithm_embeddings.get(left_doctor)
        right = algorithm_embeddings.get(right_doctor)
        if left is None or right is None:
            return 0.0
        return self._cosine_similarity(left, right)

    @staticmethod
    def _cosine_similarity(left: np.ndarray, right: np.ndarray) -> float:
        # 余弦相似度公式：
        # (v1 · v2) / (||v1|| * ||v2||)
        # 越接近 1，说明两张子图在当前算法视角下越相似。
        denominator = float(np.linalg.norm(left) * np.linalg.norm(right))
        if denominator == 0.0:
            return 0.0
        return round(float(np.dot(left, right) / denominator), 6)

    def _subgraph_export_readme(self, disease: str, doctors: list[GraphNode]) -> str:
        return (
            f"病名: {disease}\n"
            f"医家数量: {len(doctors)}\n"
            "文件说明:\n"
            "- pairwise_similarity.csv: 两两子图相似度，含 Jaccard / Graph2Vec\n"
            "- shared_nodes.csv: 多位医家共同节点\n"
            "- shared_edges.csv: 多位医家共同边\n"
            "- doctors/<医家>/core_nodes.csv: 被抽出的核心子图节点\n"
            "- doctors/<医家>/core_edges.csv: 被抽出的核心子图边\n"
            "- doctors/<医家>/audit_nodes.csv: 节点纳入理由\n"
            "- doctors/<医家>/audit_edges.csv: 边审计表\n"
            "- doctors/<医家>/relation_node_graph.json: 关系节点化后的图\n"
            "- doctors/<医家>/graph2vec_vector.csv: Graph2Vec 向量\n"
        )

    @staticmethod
    def _csv_text(rows: list[dict[str, object]], fieldnames: list[str]) -> str:
        buffer = io.StringIO()
        writer = csv.DictWriter(buffer, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fieldnames})
        return buffer.getvalue()

    @staticmethod
    def _json_text(payload: object) -> str:
        return json.dumps(payload, ensure_ascii=False, indent=2)

    @staticmethod
    def _typed_mapping(value: object) -> dict[str, object]:
        return value if isinstance(value, dict) else {}

    @staticmethod
    def _typed_list(value: object) -> list[object]:
        return value if isinstance(value, list) else []

    @staticmethod
    def _float_list(vector: np.ndarray | list[float] | tuple[float, ...]) -> list[float]:
        return [round(float(item), 8) for item in np.asarray(vector, dtype=float).tolist()]

    def _build_doctor_feature_embeddings(
        self,
        doctors: list[GraphNode],
        fastrp_payload: dict[str, object],
    ) -> list[PhysicianFeatureEmbedding]:
        embedding_map = self._typed_mapping(fastrp_payload.get("doctor_embeddings"))
        return [
            PhysicianFeatureEmbedding(
                doctor=doctor,
                patterns=self._float_list(self._vector_from_mapping(self._typed_mapping(embedding_map.get(doctor.name)), "patterns")),
                causes=self._float_list(self._vector_from_mapping(self._typed_mapping(embedding_map.get(doctor.name)), "causes")),
                mechanisms=self._float_list(self._vector_from_mapping(self._typed_mapping(embedding_map.get(doctor.name)), "mechanisms")),
                overall=self._float_list(self._vector_from_mapping(self._typed_mapping(embedding_map.get(doctor.name)), "overall")),
            )
            for doctor in doctors
        ]

    def _build_doctor_score_overview(
        self,
        doctors: list[GraphNode],
        similarity_group: PhysicianSimilarityGroup,
    ) -> list[PhysicianDoctorScore]:
        scores: dict[str, dict[str, list[float]]] = {
            doctor.name: {category: [] for category in ("patterns", "causes", "mechanisms", "overall")}
            for doctor in doctors
        }
        for category in ("patterns", "causes", "mechanisms", "overall"):
            for pair in getattr(similarity_group, category):
                scores.setdefault(pair.left_doctor, {}).setdefault(category, []).append(pair.cosine)
                scores.setdefault(pair.right_doctor, {}).setdefault(category, []).append(pair.cosine)
        return [
            PhysicianDoctorScore(
                doctor=doctor,
                scores=PhysicianCategoryScores(
                    patterns=round(self._mean(scores.get(doctor.name, {}).get("patterns", [])), 6),
                    causes=round(self._mean(scores.get(doctor.name, {}).get("causes", [])), 6),
                    mechanisms=round(self._mean(scores.get(doctor.name, {}).get("mechanisms", [])), 6),
                    overall=round(self._mean(scores.get(doctor.name, {}).get("overall", [])), 6),
                ),
            )
            for doctor in doctors
        ]

    def _build_feature_similarity_candidates(
        self,
        fastrp_payload: dict[str, object],
    ) -> list[PhysicianFeatureSimilarityCandidate]:
        return [
            PhysicianFeatureSimilarityCandidate(
                category=str(item.get("category", "")),
                left_doctor=str(item.get("left_doctor", "")),
                left_feature_name=str(item.get("left_feature_name", "")),
                right_doctor=str(item.get("right_doctor", "")),
                right_feature_name=str(item.get("right_feature_name", "")),
                similarity=round(float(item.get("similarity", 0.0)), 6),
            )
            for item in self._typed_list(fastrp_payload.get("feature_candidates"))
            if isinstance(item, dict)
        ]

    def _build_embedding_points(self, raw_points: list[object]) -> list[PhysicianEmbeddingScatterPoint]:
        items = [item for item in raw_points if isinstance(item, dict) and isinstance(item.get("vector"), list)]
        if not items:
            return []
        matrix = np.asarray([item["vector"] for item in items], dtype=float)
        projected = self._project_2d(matrix)
        return [
            PhysicianEmbeddingScatterPoint(
                id=str(item.get("id", index)),
                label=str(item.get("label", "")),
                group=str(item.get("group", "")),
                x=round(float(projected[index][0]), 6),
                y=round(float(projected[index][1]), 6),
            )
            for index, item in enumerate(items)
        ]

    def _build_doctor_embedding_points(
        self,
        embeddings: list[object],
        group: str,
        *,
        vector_attr: str = "vector",
    ) -> list[PhysicianEmbeddingScatterPoint]:
        rows = []
        for item in embeddings:
            doctor = getattr(item, "doctor", None)
            vector = getattr(item, vector_attr, [])
            if not isinstance(doctor, GraphNode) or not vector:
                continue
            rows.append({"id": doctor.id, "label": doctor.name, "group": group, "vector": vector})
        return self._build_embedding_points(rows)

    def _build_node_export_payload(
        self,
        disease: str,
        contexts: list[dict[str, object]],
        similarity_rows: dict[str, list[dict[str, object]]],
        fastrp_payload: dict[str, object],
    ) -> dict[str, object]:
        response = self.compare_nodes(disease, contexts, similarity_rows, fastrp_payload)
        return {
            "similarity_group": response.fastrp_similarity,
            "embeddings": response.doctor_feature_embeddings,
        }

    def _build_path_export_payload(self, disease: str, contexts: list[dict[str, object]]) -> dict[str, object]:
        response = self.compare_paths(disease, contexts)
        return {
            "shared_paths": response.shared_paths,
            "profiles": response.doctor_profiles,
            "similarity_pairs": response.similarity_pairs,
            "embeddings": response.embeddings,
        }

    def _similarity_export_rows(self, similarity_group: PhysicianSimilarityGroup) -> list[dict[str, object]]:
        rows: list[dict[str, object]] = []
        for category in ("patterns", "causes", "mechanisms", "overall"):
            for pair in getattr(similarity_group, category):
                rows.append(
                    {
                        "category": category,
                        "left_doctor": pair.left_doctor,
                        "right_doctor": pair.right_doctor,
                        "jaccard": pair.jaccard,
                        "overlap": pair.overlap,
                        "cosine": pair.cosine,
                    }
                )
        return rows

    def _doctor_feature_embedding_rows(self, embeddings: list[PhysicianFeatureEmbedding]) -> list[dict[str, object]]:
        rows: list[dict[str, object]] = []
        for item in embeddings:
            for category in ("patterns", "causes", "mechanisms", "overall"):
                rows.append(
                    {
                        "doctor": item.doctor.name,
                        "category": category,
                        "vector": json.dumps(getattr(item, category), ensure_ascii=False),
                    }
                )
        return rows

    def _shared_path_rows(self, shared_paths: list[SharedPhysicianPath]) -> list[dict[str, object]]:
        return [
            {
                "path_type": item.path.path_type,
                "path_text": item.path.text,
                "doctors": "、".join(item.doctors),
            }
            for item in shared_paths
        ]

    def _unique_path_rows(self, profiles: list[PhysicianPathProfile]) -> list[dict[str, object]]:
        rows: list[dict[str, object]] = []
        for profile in profiles:
            for path in profile.unique_paths:
                rows.append(
                    {
                        "doctor": profile.doctor.name,
                        "path_type": path.path_type,
                        "path_text": path.text,
                        "path_category": path.path_category,
                    }
                )
        return rows

    def _doctor_path_rows(self, profile: PhysicianPathProfile) -> list[dict[str, object]]:
        rows: list[dict[str, object]] = []
        for bucket, items in (
            ("complete", profile.complete_paths),
            ("partial", profile.partial_paths),
            ("single", profile.single_paths),
        ):
            for path in items:
                rows.append(
                    {
                        "path_category": bucket,
                        "path_type": path.path_type,
                        "path_text": path.text,
                        "signature": path.signature,
                    }
                )
        return rows

    def _path_statistics_row(self, profile: PhysicianPathProfile) -> dict[str, object]:
        return {
            "doctor": profile.doctor.name,
            "complete_count": profile.completeness.complete_count,
            "partial_count": profile.completeness.partial_count,
            "single_count": profile.completeness.single_count,
            "total_count": profile.completeness.total_count,
            "complete_ratio": profile.completeness.complete_ratio,
            "path_coverage": profile.completeness.path_coverage,
        }

    @staticmethod
    def _metapath_walk(path: PhysicianPathChain) -> list[str]:
        tokens: list[str] = []
        if path.cause:
            tokens.append(f"D:{path.cause.name}")
        if path.mechanism:
            tokens.append(f"E:{path.mechanism.name}")
        if path.pattern:
            tokens.append(f"C:{path.pattern.name}")
        if len(tokens) == 1:
            return [tokens[0], tokens[0], tokens[0]]
        if len(tokens) == 2:
            return tokens + tokens[::-1]
        if len(tokens) >= 3:
            return tokens + tokens[-2::-1]
        return []

    def _build_subgraph_embedding_profiles(
        self,
        doctors: list[GraphNode],
        embedding_payload: dict[str, object],
    ) -> list[PhysicianSubgraphEmbeddingProfile]:
        vector_embeddings = self._typed_mapping(embedding_payload.get("vector_embeddings"))
        graph2vec_embeddings = self._typed_mapping(vector_embeddings.get("Graph2Vec"))
        return [
            PhysicianSubgraphEmbeddingProfile(
                doctor=doctor,
                graph2vec_vector=self._float_list(np.asarray(graph2vec_embeddings.get(doctor.name, []), dtype=float)),
            )
            for doctor in doctors
        ]

    @staticmethod
    def _vector_from_mapping(payload: dict[str, object], key: str) -> np.ndarray:
        return np.asarray(payload.get(key, []), dtype=float)

    @staticmethod
    def _mean(values: list[float]) -> float:
        return sum(values) / len(values) if values else 0.0

    def _project_2d(self, matrix: np.ndarray) -> np.ndarray:
        if matrix.size == 0:
            return np.zeros((0, 2))
        if len(matrix) == 1:
            return np.array([[0.0, 0.0]])
        if matrix.shape[1] == 1:
            return np.column_stack((matrix[:, 0], np.zeros(len(matrix))))
        projected = PCA(n_components=2, random_state=self.RANDOM_SEED).fit_transform(matrix)
        return projected

    def _shared_node_rows(self, shared_nodes: SharedCompareNodeGroup) -> list[dict[str, object]]:
        rows: list[dict[str, object]] = []
        for category in self.SUBGRAPH_CATEGORY_ORDER:
            for item in getattr(shared_nodes, category):
                rows.append(
                    {
                        "category": category,
                        "node_name": item.node.name,
                        "node_type": item.node.type,
                        "doctors": "、".join(item.doctors),
                    }
                )
        return rows

    def _core_node_rows(
        self,
        doctor: GraphNode,
        disease: str,
        payload: dict[str, object],
    ) -> list[dict[str, object]]:
        disease_node = payload.get("disease")
        rows = [
            {
                "id": doctor.id,
                "name": doctor.name,
                "type": doctor.type,
                "role": "anchor",
                "category": "doctor",
                "inclusion_reason": "当前医家锚点",
            },
            {
                "id": disease_node.id if isinstance(disease_node, GraphNode) else "",
                "name": disease_node.name if isinstance(disease_node, GraphNode) else disease,
                "type": "B病名",
                "role": "anchor",
                "category": "disease",
                "inclusion_reason": "当前病名锚点",
            },
        ]
        node_map = payload.get("node_map", {})
        audit_nodes = payload.get("audit_nodes", [])
        if not isinstance(node_map, dict):
            return rows
        reason_by_key = {}
        for item in audit_nodes if isinstance(audit_nodes, list) else []:
            if isinstance(item, PhysicianSubgraphAuditNode):
                reason_by_key[(item.type, item.name)] = item.inclusion_reason
        for category in self.SUBGRAPH_CATEGORY_ORDER:
            category_map = node_map.get(category, {})
            if not isinstance(category_map, dict):
                continue
            for node in sorted(category_map.values(), key=lambda item: (item.type, item.name)):
                if not isinstance(node, GraphNode):
                    continue
                rows.append(
                    {
                        "id": node.id,
                        "name": node.name,
                        "type": node.type,
                        "role": "feature",
                        "category": category,
                        "inclusion_reason": reason_by_key.get((node.type, node.name), ""),
                    }
                )
        return rows

    @staticmethod
    def _serialize_vector_graph(graph: nx.Graph) -> dict[str, object]:
        return {
            "nodes": [
                {
                    "id": node_id,
                    **graph.nodes[node_id],
                }
                for node_id in sorted(graph.nodes)
            ],
            "edges": [
                {
                    "source": min(source, target),
                    "target": max(source, target),
                }
                for source, target in sorted(
                    {(min(source, target), max(source, target)) for source, target in graph.edges}
                )
            ],
        }

    @staticmethod
    def _vector_rows(vector: np.ndarray | list[float] | tuple[float, ...]) -> list[dict[str, object]]:
        values = np.asarray(vector, dtype=float).tolist()
        return [
            {
                "index": index,
                "value": round(float(value), 12),
            }
            for index, value in enumerate(values)
        ]

    @staticmethod
    def _safe_filename_component(value: str) -> str:
        safe = "".join(char if char.isalnum() or char in {"-", "_"} else "_" for char in (value or "").strip())
        return safe.strip("_") or "export"

    @staticmethod
    def _elapsed_ms(start: float) -> int:
        return int(round((time.perf_counter() - start) * 1000))

    def _report_export_readme(self, disease: str) -> str:
        return (
            f"病名: {disease}\n"
            "说明:\n"
            "- report/*.docx: 论文版 Word 报告\n"
            "- figures/*.png: 600dpi 高清论文图片\n"
            "- tables/*.csv: 图表对应原始表格\n"
            "- report_metadata.json: 导出元数据与耗时统计\n"
        )

    def _build_report_tables(
        self,
        node_result: PhysicianNodeCompareResponse,
        path_result: PhysicianPathCompareResponse,
        subgraph_result: PhysicianSubgraphCompareResponse,
    ) -> dict[str, list[dict[str, object]]]:
        return {
            "node_similarity_overall.csv": self._similarity_pair_rows(node_result.similarity.overall),
            "fastrp_similarity_overall.csv": self._similarity_pair_rows(node_result.fastrp_similarity.overall),
            "path_pairwise_similarity.csv": [
                {
                    "left_doctor": pair.left_doctor,
                    "right_doctor": pair.right_doctor,
                    "shared_path_count": pair.shared_path_count,
                    "union_path_count": pair.union_path_count,
                    "path_jaccard": pair.path_jaccard,
                    "metapath2vec_cosine": pair.metapath2vec_cosine,
                }
                for pair in path_result.similarity_pairs
            ],
            "path_completeness.csv": [
                {
                    "doctor": profile.doctor.name,
                    "complete_count": profile.completeness.complete_count,
                    "partial_count": profile.completeness.partial_count,
                    "single_count": profile.completeness.single_count,
                    "total_count": profile.completeness.total_count,
                    "complete_ratio": profile.completeness.complete_ratio,
                    "path_coverage": profile.completeness.path_coverage,
                }
                for profile in path_result.doctor_profiles
            ],
            "subgraph_pairwise_similarity.csv": [
                {
                    "left_doctor": pair.left_doctor,
                    "right_doctor": pair.right_doctor,
                    "node_jaccard": pair.node_jaccard,
                    "edge_jaccard": pair.edge_jaccard,
                    "subgraph_jaccard": pair.subgraph_jaccard,
                    "graph2vec_cosine": pair.graph2vec_cosine,
                }
                for pair in subgraph_result.similarity_pairs
            ],
        }

    @staticmethod
    def _similarity_pair_rows(pairs: list[PhysicianSimilarityPair]) -> list[dict[str, object]]:
        return [
            {
                "left_doctor": pair.left_doctor,
                "right_doctor": pair.right_doctor,
                "jaccard": pair.jaccard,
                "overlap": pair.overlap,
                "cosine": pair.cosine,
            }
            for pair in pairs
        ]

    def _build_report_figures(
        self,
        disease: str,
        node_result: PhysicianNodeCompareResponse,
        path_result: PhysicianPathCompareResponse,
        subgraph_result: PhysicianSubgraphCompareResponse,
    ) -> dict[str, bytes]:
        modules = self._require_report_modules()
        figures: dict[str, bytes] = {}
        figures["node_similarity_overall_heatmap.png"] = self._render_heatmap_figure(
            modules,
            "节点显性总体相似度",
            node_result.doctors,
            self._pair_matrix(node_result.doctors, node_result.similarity.overall, "jaccard"),
            "Jaccard",
        )
        figures["fastrp_similarity_overall_heatmap.png"] = self._render_heatmap_figure(
            modules,
            "FastRP 总体相似度",
            node_result.doctors,
            self._pair_matrix(node_result.doctors, node_result.fastrp_similarity.overall, "cosine"),
            "Cosine",
        )
        figures["fastrp_radar.png"] = self._render_radar_figure(modules, node_result)
        figures["fastrp_scatter.png"] = self._render_scatter_figure(modules, "FastRP 医家分布", node_result.embedding_points)
        figures["path_jaccard_heatmap.png"] = self._render_heatmap_figure(
            modules,
            "路径 Jaccard 相似度",
            path_result.doctors,
            self._path_pair_matrix(path_result.doctors, path_result.similarity_pairs, "path_jaccard"),
            "Path Jaccard",
        )
        figures["metapath2vec_similarity_heatmap.png"] = self._render_heatmap_figure(
            modules,
            "Metapath2Vec 路径相似度",
            path_result.doctors,
            self._path_pair_matrix(path_result.doctors, path_result.similarity_pairs, "metapath2vec_cosine"),
            "Cosine",
        )
        figures["metapath2vec_scatter.png"] = self._render_scatter_figure(modules, "Metapath2Vec 医家分布", path_result.embedding_points)
        figures["subgraph_jaccard_heatmap.png"] = self._render_heatmap_figure(
            modules,
            "子图 Jaccard 相似度",
            subgraph_result.doctors,
            self._subgraph_pair_matrix(subgraph_result.doctors, subgraph_result.similarity_pairs, "subgraph_jaccard"),
            "Subgraph Jaccard",
        )
        figures["graph2vec_similarity_heatmap.png"] = self._render_heatmap_figure(
            modules,
            "Graph2Vec 子图相似度",
            subgraph_result.doctors,
            self._subgraph_pair_matrix(subgraph_result.doctors, subgraph_result.similarity_pairs, "graph2vec_cosine"),
            "Graph2Vec",
        )
        figures["graph2vec_scatter.png"] = self._render_scatter_figure(modules, "Graph2Vec 医家分布", subgraph_result.embedding_points)
        return figures

    def _build_report_docx(
        self,
        disease: str,
        node_result: PhysicianNodeCompareResponse,
        path_result: PhysicianPathCompareResponse,
        subgraph_result: PhysicianSubgraphCompareResponse,
        figures: dict[str, bytes],
        table_payload: dict[str, list[dict[str, object]]],
        timings: dict[str, int],
    ) -> bytes:
        modules = self._require_report_modules()
        Document = modules["Document"]
        Inches = modules["Inches"]
        Pt = modules["Pt"]
        WD_ALIGN_PARAGRAPH = modules["WD_ALIGN_PARAGRAPH"]
        document = Document()
        style = document.styles["Normal"]
        style.font.size = Pt(7)
        style.font.name = "Times New Roman"
        if hasattr(style.font, "element"):
            style.font.element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体")

        document.add_heading(f"《{disease}》医家比较分析报告", level=0)
        document.add_paragraph(f"生成时间：{time.strftime('%Y-%m-%d %H:%M:%S')}")
        document.add_paragraph(f"医家数量：{node_result.doctor_count}")

        document.add_heading("一、基本信息与耗时", level=1)
        document.add_paragraph(self._report_intro_text(disease, node_result))
        self._append_docx_table(document, [
            {"项目": "分析总耗时(ms)", "数值": timings["analysis_ms"]},
            {"项目": "图片生成耗时(ms)", "数值": timings["figure_ms"]},
            {"项目": "Word 生成耗时(ms)", "数值": timings["word_ms"]},
            {"项目": "论文包总耗时(ms)", "数值": timings["total_ms"]},
        ])

        document.add_heading("二、节点比较结果", level=1)
        document.add_paragraph(self._node_report_text(node_result))
        self._append_docx_picture(document, figures["node_similarity_overall_heatmap.png"], "图1 节点显性总体相似度热力图", Inches)
        self._append_docx_picture(document, figures["fastrp_similarity_overall_heatmap.png"], "图2 FastRP 总体相似度热力图", Inches)
        self._append_docx_picture(document, figures["fastrp_radar.png"], "图3 FastRP 医家特征雷达图", Inches)
        self._append_docx_table(document, table_payload["node_similarity_overall.csv"])
        self._append_docx_table(document, table_payload["fastrp_similarity_overall.csv"])

        document.add_heading("三、辨证路径比较结果", level=1)
        document.add_paragraph(self._path_report_text(path_result))
        self._append_docx_picture(document, figures["path_jaccard_heatmap.png"], "图4 路径 Jaccard 热力图", Inches)
        self._append_docx_picture(document, figures["metapath2vec_similarity_heatmap.png"], "图5 Metapath2Vec 路径相似度热力图", Inches)
        self._append_docx_picture(document, figures["metapath2vec_scatter.png"], "图6 Metapath2Vec 医家分布图", Inches)
        self._append_docx_table(document, table_payload["path_pairwise_similarity.csv"])
        self._append_docx_table(document, table_payload["path_completeness.csv"])

        document.add_heading("四、核心子图比较结果", level=1)
        document.add_paragraph(self._subgraph_report_text(subgraph_result))
        self._append_docx_picture(document, figures["subgraph_jaccard_heatmap.png"], "图7 子图 Jaccard 热力图", Inches)
        self._append_docx_picture(document, figures["graph2vec_similarity_heatmap.png"], "图8 Graph2Vec 相似度热力图", Inches)
        self._append_docx_picture(document, figures["graph2vec_scatter.png"], "图9 Graph2Vec 医家分布图", Inches)
        self._append_docx_table(document, table_payload["subgraph_pairwise_similarity.csv"])

        document.add_heading("五、附录", level=1)
        document.add_paragraph("附件中已同步导出 figures 和 tables 目录，可直接用于论文插图和附表整理。")

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _require_report_modules(self) -> dict[str, object]:
        try:
            import matplotlib
            matplotlib.use("Agg")
            from matplotlib import font_manager, pyplot as plt
        except ModuleNotFoundError as exc:
            raise ValueError("论文图片导出依赖 matplotlib，请先在 KG 环境安装。") from exc
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt
        except ModuleNotFoundError as exc:
            raise ValueError("论文 Word 导出依赖 python-docx，请先在 KG 环境安装。") from exc
        self._assert_report_font(font_manager, "Times New Roman")
        self._assert_report_font(font_manager, "SimSun")
        return {
            "plt": plt,
            "font_manager": font_manager,
            "Document": Document,
            "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
            "Inches": Inches,
            "Pt": Pt,
        }

    @staticmethod
    def _assert_report_font(font_manager: object, font_name: str) -> None:
        try:
            font_manager.findfont(font_name, fallback_to_default=False)
        except Exception as exc:
            raise ValueError(f"论文图片导出缺少字体：{font_name}") from exc

    def _pair_matrix(
        self,
        doctors: list[GraphNode],
        pairs: list[PhysicianSimilarityPair],
        metric: str,
    ) -> np.ndarray:
        names = [doctor.name for doctor in doctors]
        matrix = np.eye(len(names), dtype=float)
        for pair in pairs:
            if pair.left_doctor in names and pair.right_doctor in names:
                left_index = names.index(pair.left_doctor)
                right_index = names.index(pair.right_doctor)
                value = float(getattr(pair, metric))
                matrix[left_index, right_index] = value
                matrix[right_index, left_index] = value
        return matrix

    def _path_pair_matrix(
        self,
        doctors: list[GraphNode],
        pairs: list[PhysicianPathSimilarityPair],
        metric: str,
    ) -> np.ndarray:
        names = [doctor.name for doctor in doctors]
        matrix = np.eye(len(names), dtype=float)
        for pair in pairs:
            if pair.left_doctor in names and pair.right_doctor in names:
                left_index = names.index(pair.left_doctor)
                right_index = names.index(pair.right_doctor)
                value = float(getattr(pair, metric))
                matrix[left_index, right_index] = value
                matrix[right_index, left_index] = value
        return matrix

    def _subgraph_pair_matrix(
        self,
        doctors: list[GraphNode],
        pairs: list[PhysicianSubgraphSimilarityPair],
        metric: str,
    ) -> np.ndarray:
        names = [doctor.name for doctor in doctors]
        matrix = np.eye(len(names), dtype=float)
        for pair in pairs:
            if pair.left_doctor in names and pair.right_doctor in names:
                left_index = names.index(pair.left_doctor)
                right_index = names.index(pair.right_doctor)
                value = float(getattr(pair, metric))
                matrix[left_index, right_index] = value
                matrix[right_index, left_index] = value
        return matrix

    def _render_heatmap_figure(
        self,
        modules: dict[str, object],
        title: str,
        doctors: list[GraphNode],
        matrix: np.ndarray,
        colorbar_label: str,
    ) -> bytes:
        plt = modules["plt"]
        font_manager = modules["font_manager"]
        zh_font = font_manager.FontProperties(family="SimSun", size=7)
        en_font = font_manager.FontProperties(family="Times New Roman", size=7)
        figure, axis = plt.subplots(figsize=(6.2, 4.8), dpi=600)
        image = axis.imshow(matrix, cmap="Blues", vmin=0.0, vmax=1.0)
        names = [doctor.name for doctor in doctors]
        axis.set_xticks(range(len(names)))
        axis.set_yticks(range(len(names)))
        axis.set_xticklabels(names, fontproperties=zh_font)
        axis.set_yticklabels(names, fontproperties=zh_font)
        axis.set_title(title, fontproperties=zh_font, pad=8)
        for spine in axis.spines.values():
            spine.set_linewidth(1.0)
        for row_index in range(len(names)):
            for column_index in range(len(names)):
                axis.text(
                    column_index,
                    row_index,
                    f"{matrix[row_index, column_index]:.3f}",
                    ha="center",
                    va="center",
                    fontproperties=en_font,
                )
        colorbar = figure.colorbar(image, ax=axis)
        colorbar.set_label(colorbar_label, fontproperties=en_font)
        return self._save_figure(figure, plt)

    def _render_scatter_figure(
        self,
        modules: dict[str, object],
        title: str,
        points: list[PhysicianEmbeddingScatterPoint],
    ) -> bytes:
        plt = modules["plt"]
        font_manager = modules["font_manager"]
        zh_font = font_manager.FontProperties(family="SimSun", size=7)
        en_font = font_manager.FontProperties(family="Times New Roman", size=7)
        figure, axis = plt.subplots(figsize=(6.2, 4.8), dpi=600)
        grouped: dict[str, list[PhysicianEmbeddingScatterPoint]] = {}
        for point in points:
            grouped.setdefault(point.group, []).append(point)
        for label, items in grouped.items():
            axis.scatter([item.x for item in items], [item.y for item in items], label=label, s=18)
            for item in items:
                axis.text(item.x, item.y, item.label, fontproperties=zh_font)
        axis.set_title(title, fontproperties=zh_font, pad=8)
        axis.set_xlabel("PCA-1", fontproperties=en_font)
        axis.set_ylabel("PCA-2", fontproperties=en_font)
        for spine in axis.spines.values():
            spine.set_linewidth(1.0)
        axis.legend(prop=zh_font, frameon=False)
        return self._save_figure(figure, plt)

    def _render_radar_figure(
        self,
        modules: dict[str, object],
        node_result: PhysicianNodeCompareResponse,
    ) -> bytes:
        plt = modules["plt"]
        font_manager = modules["font_manager"]
        zh_font = font_manager.FontProperties(family="SimSun", size=7)
        figure = plt.figure(figsize=(6.2, 4.8), dpi=600)
        axis = figure.add_subplot(111, polar=True)
        categories = ["证型", "病因", "病机", "总体"]
        angles = np.linspace(0, 2 * math.pi, len(categories), endpoint=False).tolist()
        angles += angles[:1]
        for item in node_result.similarity_overview:
            values = [
                item.scores.patterns,
                item.scores.causes,
                item.scores.mechanisms,
                item.scores.overall,
            ]
            values += values[:1]
            axis.plot(angles, values, linewidth=1.5, label=item.doctor.name)
            axis.fill(angles, values, alpha=0.08)
        axis.set_xticks(angles[:-1])
        axis.set_xticklabels(categories, fontproperties=zh_font)
        axis.set_title("FastRP 医家特征雷达图", fontproperties=zh_font, pad=12)
        axis.legend(prop=zh_font, loc="upper right", frameon=False)
        axis.spines["polar"].set_linewidth(1.0)
        return self._save_figure(figure, plt)

    @staticmethod
    def _save_figure(figure: object, plt: object) -> bytes:
        buffer = io.BytesIO()
        figure.tight_layout()
        figure.savefig(buffer, format="png", dpi=600, facecolor="white", bbox_inches="tight")
        plt.close(figure)
        return buffer.getvalue()

    @staticmethod
    def _append_docx_picture(document: object, payload: bytes, caption: str, inches_factory: object) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = 1
        run = paragraph.add_run()
        run.add_picture(io.BytesIO(payload), width=inches_factory(6.0))
        caption_paragraph = document.add_paragraph(caption)
        caption_paragraph.alignment = 1

    @staticmethod
    def _append_docx_table(document: object, rows: list[dict[str, object]]) -> None:
        normalized_rows = rows or [{"说明": "暂无数据"}]
        headers = list(normalized_rows[0].keys())
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for index, header in enumerate(headers):
            header_cells[index].text = str(header)
        for row in normalized_rows:
            cells = table.add_row().cells
            for index, header in enumerate(headers):
                cells[index].text = str(row.get(header, ""))

    def _report_intro_text(self, disease: str, node_result: PhysicianNodeCompareResponse) -> str:
        names = "、".join(doctor.name for doctor in node_result.doctors) or "暂无"
        return f"本报告围绕病名“{disease}”对 {node_result.doctor_count} 位医家开展节点、辨证路径和核心子图三层比较分析，涉及医家为：{names}。"

    def _node_report_text(self, node_result: PhysicianNodeCompareResponse) -> str:
        pair = self._top_pair(
            [
                {"label": f"{item.left_doctor} 与 {item.right_doctor}", "value": item.jaccard}
                for item in node_result.similarity.overall
            ]
        )
        fastrp_pair = self._top_pair(
            [
                {"label": f"{item.left_doctor} 与 {item.right_doctor}", "value": item.cosine}
                for item in node_result.fastrp_similarity.overall
            ]
        )
        return (
            f"节点层显性比较中，相似度最高的医家对为 {pair['label']}（{pair['value']:.3f}）；"
            f"FastRP 潜在结构相似度最高的医家对为 {fastrp_pair['label']}（{fastrp_pair['value']:.3f}）。"
        )

    def _path_report_text(self, path_result: PhysicianPathCompareResponse) -> str:
        pair = self._top_pair(
            [
                {"label": f"{item.left_doctor} 与 {item.right_doctor}", "value": item.path_jaccard}
                for item in path_result.similarity_pairs
            ]
        )
        embedding_pair = self._top_pair(
            [
                {"label": f"{item.left_doctor} 与 {item.right_doctor}", "value": item.metapath2vec_cosine}
                for item in path_result.similarity_pairs
            ]
        )
        return (
            f"路径层显性比较中，Path Jaccard 最高的医家对为 {pair['label']}（{pair['value']:.3f}）；"
            f"Metapath2Vec 语义相似度最高的医家对为 {embedding_pair['label']}（{embedding_pair['value']:.3f}）。"
        )

    def _subgraph_report_text(self, subgraph_result: PhysicianSubgraphCompareResponse) -> str:
        pair = self._top_pair(
            [
                {"label": f"{item.left_doctor} 与 {item.right_doctor}", "value": item.subgraph_jaccard}
                for item in subgraph_result.similarity_pairs
            ]
        )
        graph2vec_pair = self._top_pair(
            [
                {"label": f"{item.left_doctor} 与 {item.right_doctor}", "value": item.graph2vec_cosine}
                for item in subgraph_result.similarity_pairs
            ]
        )
        return (
            f"子图层显性比较中，子图 Jaccard 最高的医家对为 {pair['label']}（{pair['value']:.3f}）；"
            f"Graph2Vec 相似度最高的医家对为 {graph2vec_pair['label']}（{graph2vec_pair['value']:.3f}）。"
        )

    @staticmethod
    def _top_pair(rows: list[dict[str, object]]) -> dict[str, object]:
        if not rows:
            return {"label": "暂无", "value": 0.0}
        return max(rows, key=lambda item: float(item["value"]))

    def _canonical_rwr_rankings(
        self,
        graph: nx.Graph,
        doctor_id: str,
        disease_id: str,
        restart_probability: float,
    ) -> NodeRankingGroup:
        nodes = list(graph.nodes)
        if not nodes:
            return NodeRankingGroup()

        personalization = {node_id: 0.0 for node_id in nodes}
        personalization[doctor_id] = 0.5
        personalization[disease_id] = 0.5
        scores = dict(personalization)

        for _ in range(self.RWR_MAX_ITERATIONS):
            next_scores = {node_id: restart_probability * personalization[node_id] for node_id in nodes}
            for node_id in nodes:
                neighbors = list(graph.neighbors(node_id))
                if not neighbors:
                    next_scores[node_id] += (1 - restart_probability) * scores[node_id]
                    continue
                share = (1 - restart_probability) * scores[node_id] / len(neighbors)
                for neighbor_id in neighbors:
                    next_scores[neighbor_id] += share
            delta = sum(abs(next_scores[node_id] - scores[node_id]) for node_id in nodes)
            scores = next_scores
            if delta < self.RWR_TOLERANCE:
                break

        aggregated: dict[str, dict[tuple[str, str], dict[str, object]]] = {
            "patterns": {},
            "causes": {},
            "mechanisms": {},
        }
        node_by_id = self._nodes_from_graph(graph)
        for node_id, score in scores.items():
            node = node_by_id.get(node_id)
            if node is None or node.type not in self.RESULT_TYPES:
                continue
            category = self.RESULT_TYPES[node.type]
            key = self._canonical_key(node)
            payload = aggregated[category].get(
                key,
                {"node": node, "score": 0.0, "count": 0},
            )
            payload["score"] = float(payload["score"]) + float(score)
            payload["count"] = int(payload["count"]) + int(round(score * 1_000_000))
            existing = payload["node"]
            if isinstance(existing, GraphNode) and node.id < existing.id:
                payload["node"] = node
            aggregated[category][key] = payload

        return NodeRankingGroup(
            patterns=self._ranked_nodes_from_payload(aggregated["patterns"]),
            causes=self._ranked_nodes_from_payload(aggregated["causes"]),
            mechanisms=self._ranked_nodes_from_payload(aggregated["mechanisms"]),
        )

    def _ranked_nodes_from_payload(self, payload_map: dict[tuple[str, str], dict[str, object]]) -> list[RankedGraphNode]:
        ranked = [
            RankedGraphNode(
                id=payload["node"].id,
                name=payload["node"].name,
                type=payload["node"].type,
                score=round(float(payload["score"]), 6),
                count=int(payload["count"]),
            )
            for payload in payload_map.values()
            if isinstance(payload.get("node"), GraphNode)
        ]
        ranked.sort(key=lambda item: (-item.score, item.name))
        return ranked[: self.TOP_K]

    def _compare_node_group(self, grouped: dict[str, dict[tuple[str, str], GraphNode]]) -> CompareNodeGroup:
        return CompareNodeGroup(
            patterns=self._sorted_compare_nodes(grouped.get("patterns", {})),
            causes=self._sorted_compare_nodes(grouped.get("causes", {})),
            mechanisms=self._sorted_compare_nodes(grouped.get("mechanisms", {})),
        )

    @staticmethod
    def _make_subgraph_edge(
        *,
        signature: str,
        relation_type: str,
        source_name: str,
        source_type: str,
        target_name: str,
        target_type: str,
    ) -> PhysicianSubgraphEdge:
        return PhysicianSubgraphEdge(
            signature=signature,
            relation_type=relation_type,
            text=f"{source_name} -[{relation_type}]-> {target_name}",
            source_name=source_name,
            source_type=source_type,
            target_name=target_name,
            target_type=target_type,
        )

    @staticmethod
    def _make_subgraph_audit_edge(edge: PhysicianSubgraphEdge) -> PhysicianSubgraphAuditEdge:
        return PhysicianSubgraphAuditEdge(
            relation_type=edge.relation_type,
            text=edge.text,
            source_name=edge.source_name,
            source_type=edge.source_type,
            target_name=edge.target_name,
            target_type=edge.target_type,
        )

    def _relation_distribution(self, edges: list[PhysicianSubgraphEdge]) -> list[PhysicianSubgraphRelationStat]:
        counts = Counter(edge.relation_type for edge in edges)
        total = sum(counts.values())
        result = [
            PhysicianSubgraphRelationStat(
                relation_type=relation_type,
                count=count,
                ratio=round(count / total, 4) if total else 0.0,
            )
            for relation_type, count in sorted(counts.items(), key=lambda item: (-item[1], item[0]))
        ]
        return result

    @staticmethod
    def _ordered_relation_nodes(source_node: GraphNode, target_node: GraphNode, relation_type: str) -> tuple[GraphNode, GraphNode]:
        expected = relation_type.split("-")
        if len(expected) == 2:
            if source_node.type == expected[0] and target_node.type == expected[1]:
                return source_node, target_node
            if source_node.type == expected[1] and target_node.type == expected[0]:
                return target_node, source_node
        ordered = sorted((source_node, target_node), key=lambda item: (item.type, item.name, item.id))
        return ordered[0], ordered[1]

    @staticmethod
    def _subgraph_structure_signature(relation_type: str, source_node: GraphNode, target_node: GraphNode) -> str:
        return f"{relation_type}|{source_node.type}|{source_node.name}|{target_node.type}|{target_node.name}"

    @staticmethod
    def _jaccard(left: set[object], right: set[object]) -> float:
        union = left | right
        return round(len(left & right) / len(union), 6) if union else 0.0

    @staticmethod
    def _path_node_map_for_type(
        node_type: str,
        direct_patterns: dict[tuple[str, str], GraphNode],
        direct_causes: dict[tuple[str, str], GraphNode],
        direct_mechanisms: dict[tuple[str, str], GraphNode],
    ) -> dict[tuple[str, str], GraphNode]:
        if node_type == "C证型":
            return direct_patterns
        if node_type == "D病因":
            return direct_causes
        return direct_mechanisms

    def _store_canonical_node(self, node_map: dict[tuple[str, str], GraphNode], node: GraphNode) -> None:
        key = self._canonical_key(node)
        existing = node_map.get(key)
        if existing is None or node.id < existing.id:
            node_map[key] = node

    def _make_path_chain(
        self,
        *,
        path_type: str,
        path_category: str,
        cause: GraphNode | None = None,
        mechanism: GraphNode | None = None,
        pattern: GraphNode | None = None,
    ) -> PhysicianPathChain:
        return PhysicianPathChain(
            path_type=path_type,
            path_category=path_category,
            signature=self._path_signature(path_type, cause, mechanism, pattern),
            text=self._path_text(path_type, cause, mechanism, pattern),
            cause=self._compare_node(cause) if cause else None,
            mechanism=self._compare_node(mechanism) if mechanism else None,
            pattern=self._compare_node(pattern) if pattern else None,
        )

    @staticmethod
    def _path_signature(
        path_type: str,
        cause: GraphNode | None,
        mechanism: GraphNode | None,
        pattern: GraphNode | None,
    ) -> str:
        parts = [path_type]
        if cause is not None:
            parts.append(cause.name)
        if mechanism is not None:
            parts.append(mechanism.name)
        if pattern is not None:
            parts.append(pattern.name)
        return "|".join(parts)

    @staticmethod
    def _path_text(
        path_type: str,
        cause: GraphNode | None,
        mechanism: GraphNode | None,
        pattern: GraphNode | None,
    ) -> str:
        if path_type == "D-E-C":
            return f"{cause.name} -> {mechanism.name} -> {pattern.name}"
        if path_type == "D-E":
            return f"{cause.name} -> {mechanism.name}"
        if path_type == "E-C":
            return f"{mechanism.name} -> {pattern.name}"
        if path_type == "D-C":
            return f"{cause.name} -> {pattern.name}"
        if path_type == "D":
            return cause.name
        if path_type == "E":
            return mechanism.name
        return pattern.name

    def _sorted_compare_nodes(self, node_map: dict[tuple[str, str], GraphNode]) -> list[CompareNode]:
        nodes = [self._compare_node(node) for node in node_map.values()]
        nodes.sort(key=lambda item: item.name)
        return nodes

    @staticmethod
    def _compare_node(node: GraphNode) -> CompareNode:
        return CompareNode(
            id=node.id,
            name=node.name,
            label=node.label,
            type=node.type,
        )

    @staticmethod
    def _canonical_key(node: GraphNode) -> tuple[str, str]:
        return node.type, node.name

    @staticmethod
    def _canonical_key_from_compare_node(node: CompareNode) -> tuple[str, str]:
        return node.type, node.name

    def _node_compare_summary_message(self, disease: str, doctors: list[GraphNode], shared_node_count: int) -> str:
        if not doctors:
            return f"当前病名“{disease}”没有可比较的医家数据。"
        if len(doctors) == 1:
            return f"当前只有 {doctors[0].name} 一位医家，先看他的节点画像和 RWR 核心节点。"
        if shared_node_count == 0:
            return f"当前已比较 {len(doctors)} 位医家，但共同节点还不明显，先重点看各医家的独有节点。"
        return f"当前已比较 {len(doctors)} 位医家，主看 Jaccard，相同点与核心节点都已汇总在下方。"

    def _path_compare_summary_message(self, disease: str, doctors: list[GraphNode], shared_path_count: int) -> str:
        if not doctors:
            return f"当前病名“{disease}”没有可比较的辨证路径数据。"
        if len(doctors) == 1:
            return f"当前只有 {doctors[0].name} 一位医家，先看他的辨证链完整度、独有路径和路径向量。"
        if shared_path_count == 0:
            return f"当前已比较 {len(doctors)} 位医家，但共同辨证链还不明显，先重点看每位医家的独有路径和 Metapath2Vec 相似度。"
        return f"当前已比较 {len(doctors)} 位医家，完整链、Path Jaccard 和 Metapath2Vec 结果都已汇总在下方。"

    def _subgraph_compare_summary_message(
        self,
        disease: str,
        doctors: list[GraphNode],
        shared_node_count: int,
        shared_edge_count: int,
    ) -> str:
        if not doctors:
            return f"当前病名“{disease}”没有可比较的核心子图数据。"
        if len(doctors) == 1:
            return f"当前只有 {doctors[0].name} 一位医家，先核对他的核心子图和审计表。"
        if shared_node_count == 0 and shared_edge_count == 0:
            return f"当前已比较 {len(doctors)} 位医家，但共同子结构还不明显，先重点看各医家的独有节点和独有边。"
        return f"当前已比较 {len(doctors)} 位医家，主看子图Jaccard，核心共同结构和审计结果都已汇总在下方。"
